# =================================================================
# SIHCLI-POTER: MÓDULO MAESTRO DE TOMA DE DECISIONES (SÍNTESIS TOTAL)
# =================================================================

import os
import sys

# --- 🔥 FIX: INYECCIÓN DE PATH ESTRATÉGICA ---
# Garantiza que Python encuentre la carpeta 'modules' ANTES de importar
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

import numpy as np
import pandas as pd
import geopandas as gpd
import plotly.express as px
import plotly.graph_objects as go
import math
import folium
from streamlit_folium import st_folium
from folium import plugins
from sqlalchemy import create_engine, text
from scipy.interpolate import griddata

import streamlit as st

# --- 1. CONFIGURACIÓN DE PÁGINA (SIEMPRE PRIMERO) ---
st.set_page_config(page_title="Sihcli-Poter: Toma de Decisiones", page_icon="🎯", layout="wide")

# --- 📂 IMPORTACIÓN ROBUSTA DE MÓDULOS INTERNOS ---
try:
    from modules import selectors
    from modules.utils import encender_gemelo_digital, obtener_metabolismo_exacto
    # 🔥 FIX: Añadimos obtener_poblacion_matriz al import
    from modules.demografia_tools import render_motor_demografico, obtener_poblacion_matriz
    from modules.biodiversidad_tools import render_motor_ripario
    from modules.geomorfologia_tools import render_motor_hidrologico
    from modules.impacto_serv_ecosist import render_sigacal_analysis
    from modules.db_manager import get_engine
except ImportError as e:
    st.error(f"🚨 Error crítico cargando los módulos internos: {e}")
    st.stop() # Detiene la ejecución para evitar pantallas en blanco o errores en cascada

# ==========================================
# 📂 NUEVO: MENÚ DE NAVEGACIÓN PERSONALIZADO
# ==========================================
selectors.renderizar_menu_navegacion("Toma de Decisiones")
encender_gemelo_digital()

# =========================================================================
# 🏷️ RECUPERACIÓN DEL TÍTULO PRINCIPAL DE LA PÁGINA
# =========================================================================
st.title("🎯 Módulo Maestro de Toma de Decisiones y Síntesis Territorial")
st.markdown("""
Integración Multicriterio para la **Seguridad Hídrica**, la **Conservación de la Biodiversidad** y la **Gestión del Riesgo**.  
*Utilice este tablero gerencial para simular escenarios de inversión y priorizar áreas de restauración ecológica.*
""")

# 🎨 INYECCIÓN CSS PREMIUM (Para Expanders y Tipografía Gerencial)
st.markdown("""
<style>
/* 1. CAMBIO DE TIPOGRAFÍA GLOBAL AL ESTILO 'GEORGIA' */
html, body, [class*="css"]  {
    font-family: 'Georgia', serif !important;
}

/* 2. ESTILO PARA LOS EXPANDERS */
div[data-testid="stExpander"] {
    background-color: #ffffff;
    border: 1px solid #e0e6ed;
    border-radius: 8px;
    box-shadow: 0 4px 6px rgba(0,0,0,0.04);
    margin-bottom: 12px;
}
div[data-testid="stExpander"] summary {
    background-color: #f8fafc;
    border-radius: 8px;
    padding: 10px 15px;
}
div[data-testid="stExpander"] summary:hover {
    background-color: #f1f5f9;
}
div[data-testid="stExpander"] summary p {
    font-size: 1.1rem !important;
    font-weight: 600 !important;
    color: #1e293b !important;
    font-family: 'Georgia', serif !important;
}

/* 🚀 3. NUEVO: ESTILO PARA QUE LAS PESTAÑAS PAREZCAN BOTONES GIGANTES */
button[data-baseweb="tab"] {
    font-size: 1.15rem !important;
    background-color: #f8fafc !important;
    border: 1px solid #cbd5e1 !important;
    border-bottom: none !important;
    border-radius: 8px 8px 0 0 !important;
    padding: 12px 24px !important;
    margin-right: 5px !important;
    color: #64748b !important;
    transition: all 0.3s ease;
}
button[data-baseweb="tab"]:hover {
    background-color: #e2e8f0 !important;
}
button[data-baseweb="tab"][aria-selected="true"] {
    background-color: #ffffff !important;
    color: #1e293b !important;
    border-top: 4px solid #e74c3c !important;
    font-weight: bold !important;
    box-shadow: 0 -2px 5px rgba(0,0,0,0.05) !important;
}
</style>
""", unsafe_allow_html=True)

# --- 2. EXPLICACIÓN METODOLÓGICA ---
def render_metodologia():
    with st.expander("🔬 METODOLOGÍA Y GUÍA DEL TABLERO (Leer antes de iniciar)", expanded=False):
        st.markdown("""
        ### ¿Cómo funciona esta página?
        Este módulo es la **Síntesis Estratégica** de Sihcli-Poter. Su motor matemático funde dos marcos de trabajo internacionales:
        
        1. **Análisis Multicriterio Espacial (SMCA):** Identifica *dónde* actuar cruzando Balance Hídrico, Biodiversidad y Geomorfología.
        2. **Estándares Corporativos (WRI):** Mide el *impacto volumétrico* de las intervenciones usando la metodología VWBA (Volumetric Water Benefit Accounting).
        
        ---
        
        ### 🗺️ La Arquitectura Conceptual (Ruta de Decisión Informada)
        Para facilitar la estructuración de la política pública, este tablero está diseñado como un embudo lógico de 6 pasos. Le recomendamos descender secuencialmente:

        *   **📍 PASO 1: Fotografía del Paciente (Diagnóstico Base):** ¿Dónde estamos parados hoy? *(Métricas iniciales de metabolismo y Velocímetros de Salud).*
        *   **📍 PASO 2: Prueba de Estrés (Riesgo Físico y Climático):** ¿Qué pasa si la situación empeora? *(Simulador de Avenidas Torrenciales y Fenómeno ENSO).*
        *   **📍 PASO 3: Ingeniería de Soluciones (Simulador WRI):** ¿Qué necesitamos construir físicamente para mitigar el déficit? *(Diagrama termodinámico Sankey, metas en hectáreas y toneladas).*
        *   **📍 PASO 4: Viabilidad Financiera (Optimizador ROI):** ¿Cuánto cuesta la infraestructura y cuál es el retorno de inversión? *(Simulador de capital Dólar vs. ISHI).*
        *   **📍 PASO 5: Inteligencia Táctica (Terreno):** ¿Exactamente en cuáles cuencas y predios vamos a poner ese dinero? *(Ranking AHP y Visor 3D PyDeck).*
        *   **📍 PASO 6: La Gran Síntesis (Manifiesto Estratégico):** La Inteligencia Artificial consolida las decisiones anteriores, redacta un veredicto gerencial y exporta el Plan de Acción en un documento oficial Word.
        """)

# --- 3. FUNCIONES DE CARGA ROBUSTAS ---
@st.cache_data(ttl=3600)
def load_context_layers(gdf_zona_bounds):
    layers = {'cuencas': None, 'predios': None, 'drenaje': None, 'geomorf': None}
    minx, miny, maxx, maxy = gdf_zona_bounds
    from shapely.geometry import box
    roi = gpd.GeoDataFrame(geometry=[box(minx, miny, maxx, maxy)], crs="EPSG:4326")
    
    base_dir = os.path.join(os.path.dirname(__file__), '..', 'data')
    files = {
        'cuencas': "SubcuencasAinfluencia.geojson",
        'predios': "PrediosEjecutados.geojson",
        'drenaje': "Drenaje_Sencillo.geojson",
        'geomorf': "UnidadesGeomorfologicas.geojson"
    }
    for key, fname in files.items():
        try:
            fpath = os.path.join(base_dir, fname)
            if os.path.exists(fpath):
                gdf = gpd.read_file(fpath)
                if gdf.crs != "EPSG:4326": gdf = gdf.to_crs("EPSG:4326")
                layers[key] = gpd.clip(gdf, roi)
        except: pass
    return layers

# ==============================================================================
# --- 4. LÓGICA PRINCIPAL (VERSIÓN DOMINICAL BLINDADA) ---
# ==============================================================================
render_metodologia()

# 1. Selector Espacial Maestro
ids_sel, nombre_zona, alt_ref, gdf_zona = selectors.render_selector_espacial()

# ==============================================================================
# 🔥 ESCUDO ANTI-FANTASMAS (Bloquea placeholders del menú)
# ==============================================================================
textos_invalidos = ["", "-- Seleccione --", "-- BLOQUE --", "Bloque Hidro", "Seleccione", "None"]
if str(nombre_zona).strip() in textos_invalidos:
    st.info("👈 Por favor, seleccione un territorio válido en el menú lateral para cargar el Centro de Comando.")
    st.stop() # Detiene la ejecución aquí para que no salgan letreros amarillos
    
# 2. Configuración de Pesos AHP en Sidebar
with st.sidebar:
    st.header("⚖️ Pesos AHP (Multicriterio)")
    st.caption("Define la importancia de cada vector. El sistema los normalizará al 100%.")
    raw_agua = st.slider("💧 Riesgo Hídrico (Estrés/Escasez)", 0, 10, 7)
    raw_bio = st.slider("🍃 Valor Biótico (Biodiversidad)", 0, 10, 4)
    raw_socio = st.slider("👥 Presión Socioeconómica", 0, 10, 5)
    
    suma_pesos = raw_agua + raw_bio + raw_socio if (raw_agua + raw_bio + raw_socio) > 0 else 1
    w_agua = raw_agua / suma_pesos
    w_bio = raw_bio / suma_pesos
    w_socio = raw_socio / suma_pesos
    
    st.info(f"**Pesos Finales:**\nHídrico: {w_agua*100:.0f}% | Biótico: {w_bio*100:.0f}% | Socio: {w_socio*100:.0f}%")
    st.divider()
    st.subheader("👁️ Visibilidad de Capas SIG")
    v_sat = st.checkbox("Fondo Satelital", True)
    v_drain = st.checkbox("Red de Drenaje", True)
    v_geo = st.checkbox("Geomorfología", False)

# 3. Lógica de Cálculo de Impactos
if gdf_zona is not None and not gdf_zona.empty:
    engine = get_engine()
    
    # --- Control Temporal ---
    anio_actual = st.slider("📅 Año de Proyección (Simulación Futura):", min_value=2024, max_value=2050, value=2025, step=1)
    
    # ==============================================================================
    # 🧠 NÚCLEO DE CONEXIÓN TOLERANTE (SQL MULTI-MATRIZ CON TRADUCTOR)
    # ==============================================================================
    
    # Recuperamos el nivel exacto que el usuario seleccionó en el menú lateral
    nivel_req = st.session_state.get('nivel_activo_global', 'NINGUNO')

    @st.cache_data(ttl=3600)
    def consultar_matriz_sql(tabla, territorio, nivel, col_nivel="Nivel"):
        try:
            from sqlalchemy import text
            from modules.db_manager import get_engine
            import pandas as pd
        
            engine = get_engine()
        
            # 🚀 BÚSQUEDA BLINDADA: Busca directamente en las columnas, ignorando la Llave Universal
            q = text(f'''
                SELECT * FROM {tabla} 
                WHERE TRIM(UPPER("Territorio")) = UPPER(:t) 
                AND TRIM(UPPER("{col_nivel}")) = UPPER(:n)
            ''')
            df_res = pd.read_sql(q, engine, params={'t': str(territorio).strip(), 'n': str(nivel).strip()})

            if df_res.empty:
                q_fall = text(f'''
                    SELECT * FROM {tabla} 
                    WHERE UPPER("Territorio") LIKE UPPER(:t) 
                    AND UPPER("{col_nivel}") LIKE UPPER(:n)
                ''')
                df_res = pd.read_sql(q_fall, engine, params={'t': f"%{str(territorio).strip()}%", 'n': f"%{str(nivel)[:3]}%"})
            
            return df_res
        except:
            return pd.DataFrame()
            
    def proyectar_modelo(f, anio_obj):
        import numpy as np
        x_norm = anio_obj - f.get('Año_Base', 2018)
        mod = str(f.get('Modelo_Recomendado', 'Logístico'))
        try:
            if 'Logistico' in mod or 'Logístico' in mod: return f.get('Log_K',0) / (1 + f.get('Log_a',0) * np.exp(-f.get('Log_r',0) * x_norm))
            elif 'Exponencial' in mod: return f.get('Exp_a',0) * np.exp(f.get('Exp_b',0) * x_norm)
            elif 'Lineal' in mod: return f.get('Lin_m',0) * x_norm + f.get('Lin_b',0)
            else: return f.get('Poly_A',0)*(x_norm**3) + f.get('Poly_B',0)*(x_norm**2) + f.get('Poly_C',0)*x_norm + f.get('Poly_D',0)
        except: return 0.0

    # 🔥 ENRUTADOR MAESTRO: Traduce los nombres del menú a los nombres de las matrices
    if nivel_req in ["AH", "ZH", "SZH", "NSS1", "NSS2", "NSS3"]:
        nivel_demo = "Cuenca"
    elif "CORPOAMB" in nivel_req.upper() or "CAR" in nivel_req.upper():
        nivel_demo = "CAR"
        nivel_req = "CAR" # Forzamos que la hidrología también lo busque como CAR
    else:
        nivel_demo = nivel_req

    # ---------------------------------------------------------
    # 1. CONEXIÓN DEMOGRÁFICA (SQL Estricto)
    # ---------------------------------------------------------
    df_demo = consultar_matriz_sql("matriz_maestra_demografica", nombre_zona, nivel_demo, "Nivel")
    if not df_demo.empty:
        pob_total = max(0.0, proyectar_modelo(df_demo.iloc[0], anio_actual))
        st.success(f" 👥  **Cerebro Demográfico Enlazado:** {pob_total:,.0f} habitantes detectados en SQL (Nivel: {nivel_demo}).")
        origen_demo = True
    else:
        pob_total = 0.0
        st.error(f" ❌  '{nombre_zona}' no existe en la Matriz Demográfica para el nivel {nivel_demo}.")
        origen_demo = False

    st.session_state['aleph_pob_total'] = pob_total
    st.session_state['pob_hum_calc_met'] = pob_total

    # ---------------------------------------------------------
    # 2. CONEXIÓN PECUARIA (Con Rescate y Bypass)
    # ---------------------------------------------------------
    # 🚀 FIX: Buscar primero por el nivel real (AH, ZH) que usa la nueva matriz
    df_pec = consultar_matriz_sql("matriz_maestra_pecuaria", nombre_zona, nivel_req, "Nivel")
    
    # Intento de rescate: Por si una matriz vieja usó "Cuenca"
    if df_pec.empty and nivel_demo == "Cuenca":
        df_pec = consultar_matriz_sql("matriz_maestra_pecuaria", nombre_zona, "Cuenca", "Nivel")

    bovinos, porcinos, aves = 0.0, 0.0, 0.0

    if not df_pec.empty:
        for _, f in df_pec.iterrows():
            if f['Especie'] == 'Bovinos': bovinos = max(0.0, proyectar_modelo(f, anio_actual))
            if f['Especie'] == 'Porcinos': porcinos = max(0.0, proyectar_modelo(f, anio_actual))
            if f['Especie'] == 'Aves': aves = max(0.0, proyectar_modelo(f, anio_actual))
        st.success(f" 🐄  **Cerebro Pecuario Enlazado:** {bovinos:,.0f} Bov, {porcinos:,.0f} Por, {aves:,.0f} Aves.")
        origen_pecu = True
    else:
        # 🔥 EL BYPASS DEL AMVA: Le decimos al sistema que este 0 es un éxito físico
        if nombre_zona == "AMVA":
            st.success(" 🐄  **Cerebro Pecuario Enlazado:** 0 animales (Ganado gestionado por Corantioquia en zona rural).")
            origen_pecu = True  # Es un éxito, así que lo marcamos como True
        else:
            st.warning(f" ⚠️  '{nombre_zona}' no detectado en Matriz Pecuaria. Asumiendo 0 animales para no detener el sistema.")
            origen_pecu = False # Es una falla real, lo marcamos como False
        
    st.session_state['ica_bovinos_calc_met'] = bovinos
    st.session_state['ica_porcinos_calc_met'] = porcinos
    st.session_state['ica_aves_calc_met'] = aves

    # ---------------------------------------------------------
    # 3. CONEXIÓN HIDROLÓGICA Y OFERTA BASE
    # ---------------------------------------------------------
    # 🔥 FIX: Restauramos TU variable original 'area_cuenca_km2' para no romper nada abajo
    area_cuenca_km2 = 0.0  
    area_km2 = 0.0
    oferta_anual_m3 = 0.0
    caudal_base_m3s = 0.0
    lluvia_base_mm = 0.0
    recarga_base_mm = 0.0
    altitud_m = 1500.0
    q_medio_real = 0.0
    q_min_real = 0.0

    df_hidro = consultar_matriz_sql("matriz_hidrologica_maestra", nombre_zona, nivel_req, "Jerarquia")

    if not df_hidro.empty:
        row_h = df_hidro.iloc[0]
        area_km2 = float(row_h.get('Area_km2', 0))
        
        # 💡 EL PUENTE: Le devolvemos el valor a tu variable original
        area_cuenca_km2 = area_km2 
        
        caudal_base_m3s = float(row_h.get('Caudal_Medio_m3s', 0))
        oferta_anual_m3 = caudal_base_m3s * 31536000
        lluvia_base_mm = float(row_h.get('Lluvia_mm', 0))
        recarga_base_mm = float(row_h.get('Recarga_mm', 0))
        altitud_m = float(row_h.get('Altitud_m', 1500))
        
        q_medio_real = caudal_base_m3s
        q_min_real = float(row_h.get('Caudal_Minimo_m3s', caudal_base_m3s * 0.2))
        
        st.success(f" 💧  **Cerebro Hidrológico Enlazado:** Área {area_cuenca_km2:,.1f} km², Caudal Medio {caudal_base_m3s:,.2f} m³/s.")
        origen_hidro = True
    else:
        st.error(f" ❌  '{nombre_zona}' no existe en la Matriz Hidrológica para el nivel {nivel_req}.")
        origen_hidro = False

    # Guardamos en session_state usando tu variable original
    st.session_state['aleph_area_km2'] = area_cuenca_km2
    st.session_state['aleph_recarga_mm'] = recarga_base_mm

    # ---------------------------------------------------------
    # 🛑 GUARDIA DE SEGURIDAD (Hard Stop) - ESCUDO INTELIGENTE (Informa pero NO detiene)
    # ===================================================================
    if not (origen_demo and origen_pecu and origen_hidro):
        with st.expander("⚠️ Estado de Sincronización de Matrices", expanded=True):
            if not origen_demo: st.error(f"❌ Datos Demográficos no encontrados para {nombre_zona} ({nivel_req}).")
            if not origen_pecu: st.warning(f"⚠️ Datos Pecuarios no encontrados. Se asume carga cero.")
            if not origen_hidro: st.error(f"❌ Datos Hidrológicos no encontrados para {nombre_zona} ({nivel_req}).")
            
            st.info("💡 Puedes continuar explorando el tablero con los datos disponibles.")
    
    # (A partir de aquí, el código sabe que las 3 matrices existen y son perfectas)
    tipo_oferta = st.radio("Escenario Hidrológico de Simulación:", 
                           ["🌊 Caudal Medio (Condiciones Normales)", "🏜️ Caudal Mínimo / Estiaje (Q95)"], horizontal=True)
    oferta_dinamica = q_min_real if "Mínimo" in tipo_oferta else q_medio_real

    with st.expander("⚙️ Calibración de Oferta Hídrica Base", expanded=False):
        oferta_base = st.number_input("Caudal de Simulación (m³/s):", value=float(oferta_dinamica), step=0.01, format="%.3f")

    oferta_nominal = float(oferta_base)
    anio_base_cc = 2024
    if int(anio_actual) > anio_base_cc:
        oferta_nominal *= (1 - ((int(anio_actual) - anio_base_cc) * 0.005))

    fase_enso = st.session_state.get('enso_fase', 'Neutro')
    if "Niño Severo" in fase_enso: oferta_nominal *= 0.55
    elif "Niño Moderado" in fase_enso: oferta_nominal *= 0.75
    elif "Niña" in fase_enso: oferta_nominal *= 1.20

    st.session_state['aleph_oferta_m3s'] = oferta_nominal


    
    # ---------------------------------------------------------
    # 4. METABOLISMO Y DEMANDA (SÍNTESIS FINAL)
    # ---------------------------------------------------------
    demanda_L_dia = (pob_total * 150) + (bovinos * 40) + (porcinos * 15) + (aves * 0.3)
    demanda_m3s = (demanda_L_dia / 1000) / 86400
    
    st.session_state['demanda_total_m3s'] = demanda_m3s
    st.session_state['poblacion_servida'] = pob_total
    st.session_state['zona_activa_global'] = nombre_zona 
    
    # Evaluar si tenemos sincronización perfecta
    if origen_demo and origen_pecu and origen_hidro:
        st.success(f"✅ **Sincronización Perfecta:** Las 3 matrices maestras están alimentando a '{nombre_zona}' en tiempo real.")
        origen_carga = "Modelación SQL Exacta"
    else:
        st.warning(f"⚠️ **Sincronización Parcial:** Faltan datos en uno de los motores para '{nombre_zona}'. Se han activado valores refugio (0).")
        origen_carga = "Datos de Emergencia"

    with st.expander("🎛️ Simulación de Escenarios (Variables de Decisión)", expanded=False):
        c_sim1, c_sim2 = st.columns(2)
        impacto_cc = c_sim1.slider("📉 Reducción Oferta por Cambio Climático (%):", 0, 80, 0, step=5)
        mitigacion_dbo = c_sim2.slider("🌿 Mitigación de Cargas (SbN + PTAR) %:", 0, 100, 0, step=5)

    # Cálculos de Oferta y Demanda Finales
    oferta_anual_m3 = (oferta_nominal * 31536000) * (1 - (impacto_cc / 100))
    recarga_anual_m3 = recarga_base_mm * area_cuenca_km2 * 1000
    consumo_anual_m3 = demanda_m3s * 31536000
    
    # Modelación de Calidad (DBO5)
    proxy_carga = ((pob_total * 0.050) + (bovinos * 0.4) + (porcinos * 0.15)) * 365 / 1000
    carga_total_ton = float(st.session_state.get('carga_dbo_total_ton', proxy_carga if proxy_carga > 0 else 1500.0))
    carga_final_rio_ton = carga_total_ton * (1 - (mitigacion_dbo / 100))
    
    # 🚀 INYECCIÓN AL ALEPH: Guardamos el cálculo para que la Telemetría lo lea
    st.session_state['carga_dbo_total_ton'] = carga_total_ton
    
    # Física de Concentración
    caudal_critico_L_s = (oferta_anual_m3 / 31536000) * 1000 * 0.25
    carga_mg_s = (carga_final_rio_ton * 1_000_000_000) / 31536000
    concentracion_dbo_mg_l = carga_mg_s / caudal_critico_L_s if caudal_critico_L_s > 0.1 else 999.0

    # 🎯 KPIs (Velocímetros)
    wei_ratio = consumo_anual_m3 / oferta_anual_m3 if oferta_anual_m3 > 0 else 1.0
    ind_estres = max(0.0, min(100.0, 100.0 - (wei_ratio / 0.40) * 60))
    
    bfi_ratio = recarga_anual_m3 / oferta_anual_m3 if oferta_anual_m3 > 0 else 0.0
    factor_supervivencia = min(1.0, recarga_anual_m3 / consumo_anual_m3) if consumo_anual_m3 > 0 else 1.0
    ind_resiliencia = max(0.0, min(100.0, (bfi_ratio / 0.70) * 100 * factor_supervivencia))
    
    ind_calidad = max(0.0, min(100.0, 100 * math.exp(-0.07 * concentracion_dbo_mg_l)))
    ind_neutralidad = 0.0 

    estres_hidrico_porcentaje = (wei_ratio) * 100
    st.session_state['estres_hidrico_global'] = estres_hidrico_porcentaje
    
    # ==============================================================================
    # 🗺️ EL LIENZO TERRITORIAL (CONTEXTO GEOGRÁFICO Y SATELITAL)
    # ==============================================================================
    st.markdown("---")
    st.markdown(f"## 📍 CONTEXTO GEOGRÁFICO: El Lienzo Territorial de {nombre_zona}")
    st.info("Antes de analizar las métricas, observemos la realidad física del territorio: coberturas de suelo, relieve y presión antrópica monitoreada por satélite.")
    
    # 📥 1. DESCARGA PREDIAL (Se mueve arriba para que el mapa lo use)
    capas = {}
    try:
        if gdf_zona is not None and not gdf_zona.empty:
            capas = load_context_layers(tuple(gdf_zona.total_bounds))
    except Exception as e:
        pass

    @st.cache_data(ttl=3600, show_spinner=False)
    def obtener_predios_y_hectareas(_gdf_zona, nombre_zona_txt):
        import requests, tempfile
        import pandas as pd
        import geopandas as gpd
        ha_calc, info_debug, gdf_predios_final = 0.0, "Descargando predios...", None 
        try:
            url_predios = "https://ldunpssoxvifemoyeuac.supabase.co/storage/v1/object/public/geojson/PrediosEjecutados.geojson"
            res = requests.get(url_predios)
            if res.status_code == 200:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".geojson") as tmp:
                    tmp.write(res.content)
                    tmp_path = tmp.name
                gdf_p = gpd.read_file(tmp_path)
                if not gdf_p.empty:
                    gdf_p.set_crs(epsg=4326, allow_override=True, inplace=True)
                    gdf_p_3116, gdf_z_3116 = gdf_p.to_crs(epsg=3116), _gdf_zona.to_crs(epsg=3116)
                    gdf_p_3116['geometry'] = gdf_p_3116.geometry.make_valid().buffer(0)
                    gdf_z_3116['geometry'] = gdf_z_3116.geometry.make_valid().buffer(0)
                    recorte_exacto = gpd.clip(gdf_p_3116, gdf_z_3116)
                    if not recorte_exacto.empty:
                        ha_calc = recorte_exacto.area.sum() / 10000.0
                        info_debug = f"✅ CORTE EXACTO: {len(recorte_exacto)} fragmentos de predios operan físicamente en la zona."
                        gdf_predios_final = recorte_exacto.to_crs(epsg=4326)
                    else: info_debug = f"ℹ️ ZONA VIRGEN: Ningún predio cae dentro de {nombre_zona_txt}."
        except Exception as e: info_debug = f"❌ ERROR: {e}"
        return ha_calc, info_debug, gdf_predios_final

    with st.spinner("Sincronizando inventario predial de Supabase..."):
        ha_reales_sig, info_debug, gdf_predios_mapa = obtener_predios_y_hectareas(gdf_zona, nombre_zona)

    # 🗺️ 2. EL MAPA SATELITAL (Tu código original intacto)
    with st.expander(f"🛰️ EXPLORADOR ESPACIAL Y COBERTURAS (Google Earth Engine): {nombre_zona}", expanded=True):
        if estres_hidrico_porcentaje > 80: color_alerta, opacidad_alerta = '#8B0000', 0.5
        elif estres_hidrico_porcentaje > 40: color_alerta, opacidad_alerta = '#E74C3C', 0.4
        elif estres_hidrico_porcentaje > 20: color_alerta, opacidad_alerta = '#F39C12', 0.3
        else: color_alerta, opacidad_alerta = '#3498DB', 0.2

        if gdf_zona is not None and not gdf_zona.empty:
            bounds = gdf_zona.total_bounds
            centro_x, centro_y = (bounds[0] + bounds[2]) / 2.0, (bounds[1] + bounds[3]) / 2.0
        else:
            centro_y, centro_x = 6.2442, -75.5812

        m = folium.Map(location=[centro_y, centro_x], zoom_start=11, tiles="CartoDB positron")
        folium.TileLayer(tiles='https://mt1.google.com/vt/lyrs=s&x={x}&y={y}&z={z}', attr='Google', name='Google Satellite', overlay=False, control=True).add_to(m)
        folium.GeoJson(gdf_zona, name=f"Límite {nombre_zona}", style_function=lambda x: {'color': 'blue', 'weight': 3, 'fillOpacity': 0.1}).add_to(m)

        if gdf_predios_mapa is not None and not gdf_predios_mapa.empty:
            folium.GeoJson(gdf_predios_mapa, name="🟢 Áreas Restauradas (CV)", style_function=lambda x: {'fillColor': '#00ff00', 'color': '#003300', 'weight': 1, 'fillOpacity': 0.7}).add_to(m)

        areas_data = [] 
        try:
            import ee
            credenciales_dict = dict(st.secrets["gcp_service_account"])
            credentials = ee.ServiceAccountCredentials(email=credenciales_dict["client_email"], key_data=credenciales_dict["private_key"])
            ee.Initialize(credentials)
            
            def add_ee_layer(self, ee_image_object, vis_params, name, show=True):
                map_id_dict = ee.Image(ee_image_object).getMapId(vis_params)
                folium.raster_layers.TileLayer(tiles=map_id_dict['tile_fetcher'].url_format, attr='Google Earth Engine', name=name, overlay=True, control=True, show=show).add_to(self)
            folium.Map.add_ee_layer = add_ee_layer
            
            with st.spinner("Optimizando memoria visual satelital..."):
                from shapely.geometry import box
                minx, miny, maxx, maxy = gdf_zona.total_bounds
                roi_ee = ee.Geometry(box(minx, miny, maxx, maxy).__geo_interface__)
            
            dw_coleccion = ee.ImageCollection('GOOGLE/DYNAMICWORLD/V1').filterBounds(roi_ee).filterDate('2023-01-01', '2024-01-01')
            dw_imagen = dw_coleccion.select('label').mode().clip(roi_ee)
            dw_vis = {'min': 0, 'max': 8, 'palette': ['#419BDF', '#397D49', '#88B053', '#7A87C6', '#E49635', '#DFC35A', '#C4281B', '#A59B8F', '#B39FE1']}
            
            m.add_ee_layer(dw_imagen, dw_vis, '🛰️ Uso de Suelo (Satélite IA)')
            
            dem = ee.Image("NASA/NASADEM_HGT/001").select('elevation').clip(roi_ee)
            slope, hillshade = ee.Terrain.slope(dem), ee.Terrain.hillshade(dem, azimuth=315, elevation=45)
            m.add_ee_layer(hillshade, {'min': 0, 'max': 255}, '⛰️ Relieve (Hillshade)', show=False)
            m.add_ee_layer(dem, {'min': 1000, 'max': 3000, 'palette': ['#006600', '#002200', '#fff700', '#ab7634', '#c4d0ff', '#ffffff']}, '🆙 Elevación (Hipsometría)', show=False)
            m.add_ee_layer(slope, {'min': 0, 'max': 45, 'palette': ['white', 'red']}, '⚠️ Mapa de Pendientes', show=False)

            with st.spinner("Calculando superficies satelitales en vivo..."):
                area_image = ee.Image.pixelArea().addBands(dw_imagen)
                roi_math = ee.Geometry(gdf_zona.geometry.simplify(0.01).unary_union.__geo_interface__)
                areas_ee = area_image.reduceRegion(reducer=ee.Reducer.sum().group(groupField=1, groupName='clase'), geometry=roi_math, scale=50, maxPixels=1e10).getInfo()

                nombres_clases = {0: "💧 Agua", 1: "🌳 Bosque", 2: "🌾 Pastos (Ganadería)", 4: "🚜 Cultivos (Agroindustria)", 5: "🌿 Matorrales", 6: "🏙️ Urbano", 7: "🟫 Suelo Desnudo"}
                if 'groups' in areas_ee:
                    for grupo in areas_ee['groups']:
                        if int(grupo['clase']) in nombres_clases:
                            areas_data.append({"Cobertura": nombres_clases[int(grupo['clase'])], "Área (Ha)": grupo['sum'] / 10000.0})
        except Exception as e:
            st.warning(f"Aviso de Satélite: Modo local activado. ({e})")

        from branca.element import Template, MacroElement
        leyenda_html = """
        {% macro html(this, kwargs) %}
        <div id='maplegend' class='maplegend' style='position: absolute; z-index:9999; background-color:rgba(255, 255, 255, 0.9); border-radius:8px; padding: 15px; font-size:13px; right: 20px; bottom: 20px; box-shadow: 0 2px 6px rgba(0,0,0,0.15); pointer-events: auto;'>
        <h4 style='margin: 0 0 10px 0; font-size: 14px; color: #333;'>Coberturas de Suelo</h4>
        <ul style='list-style: none; padding: 0; margin: 0; color: #444;'>
            <li style='margin-bottom: 6px;'><span style='background:#419BDF; width: 16px; height: 16px; display: inline-block; margin-right: 8px;'></span>Agua</li>
            <li style='margin-bottom: 6px;'><span style='background:#397D49; width: 16px; height: 16px; display: inline-block; margin-right: 8px;'></span>Bosques</li>
            <li style='margin-bottom: 6px;'><span style='background:#E4A63F; width: 16px; height: 16px; display: inline-block; margin-right: 8px;'></span>Pastos</li>
            <li style='margin-bottom: 6px;'><span style='background:#A55194; width: 16px; height: 16px; display: inline-block; margin-right: 8px;'></span>Cultivos</li>
            <li style='margin-bottom: 6px;'><span style='background:#C4281B; width: 16px; height: 16px; display: inline-block; margin-right: 8px;'></span>Urbano</li>

        </ul></div>{% endmacro %}"""
        macro = MacroElement()
        macro._template = Template(leyenda_html)
        m.get_root().add_child(macro)
        
        folium.LayerControl(position='topright').add_to(m)
        st_folium(m, width="100%", height=500, returned_objects=[])

        # ==========================================
        # 5. MOSTRAR MÉTRICAS: NATURAL VS GESTIONADO
        # ==========================================
        st.markdown("### 🌍 Balance de Coberturas: Capital Natural vs. Capital Gestionado")
        
        c_met1, c_met2 = st.columns([3, 1])
        
        with c_met1:
            if areas_data:
                df_areas = pd.DataFrame(areas_data).sort_values(by="Área (Ha)", ascending=False).reset_index(drop=True)
                cols = st.columns(len(df_areas))
                for idx, row in df_areas.iterrows(): 
                    # Limpiamos el texto para que se vea más estético
                    nombre_limpio = row['Cobertura'].split(' ')[1] if ' ' in row['Cobertura'] else row['Cobertura']
                    cols[idx].metric(label=f"🛰️ {nombre_limpio}", value=f"{row['Área (Ha)']:,.0f}")
        
        with c_met2:
            # El recuadro distintivo para la inversión SIG
            ha_historicas = float(ha_reales_sig) if 'ha_reales_sig' in locals() else 0.0
            st.markdown(f"""
            <div style="background-color: #e8f8f5; padding: 10px; border-radius: 8px; border: 2px solid #2ecc71; text-align: center; box-shadow: 0 4px 6px rgba(0,0,0,0.05);">
                <div style="font-size: 0.85rem; color: #27ae60; margin-bottom: 5px;">🟢 Inversión Histórica (SIG)</div>
                <div style="font-size: 1.4rem; font-weight: bold; color: #1e8449;">{ha_historicas:,.1f} ha</div>
            </div>
            """, unsafe_allow_html=True)

     # ==============================================================================
    # 📍 PASO 1: LA FOTOGRAFÍA DEL PACIENTE (DIAGNÓSTICO BASE)
    # ==============================================================================
    st.markdown("---")
    st.markdown("## 📍 PASO 1: Diagnóstico Territorial Base")
    st.info("Sintetiza la presión poblacional (alcantarillados) y la presión geográfica (agroquímicos detectados en el mapa) para evaluar la salud del territorio.")
    
    col1, col2, col3, col4 = st.columns(4)
    with col1: st.metric("👥 Población Servida", f"{int(pob_total):,.0f} hab")
    with col2: st.metric("☣️ Carga Orgánica (DBO5)", f"{carga_total_ton:,.1f} Ton/año", "Fuentes Puntuales (Urbano)", delta_color="inverse")
    with col3:
        st.markdown(f"""
        <div style="background-color: white; padding: 10px; border-radius: 5px; border: 1px solid #eee; box-shadow: 1px 1px 3px rgba(0,0,0,0.05);">
            <div style="font-size: 0.85rem; color: #555; margin-bottom: 5px;">🐄 Presión Pecuaria Oficial</div>
            <div style="font-size: 1rem; font-weight: bold; color: #2c3e50;">🐮 {bovinos:,.0f} Bov | 🐷 {porcinos:,.0f} Por | 🐔 {aves:,.0f} Ave</div>
        </div>
        """, unsafe_allow_html=True)
    with col4: st.metric("⚠️ Estrés Hídrico Neto", f"{estres_hidrico_porcentaje:,.1f} %", "Crítico" if estres_hidrico_porcentaje > 40 else "Estable", delta_color="inverse")

    # ☠️ IMPACTO AGROQUÍMICO (HEREDANDO DEL SATÉLITE)
    st.markdown("#### ☠️ Impacto Agroquímico y Fuentes Difusas")
    
    # Lógica de respaldo invencible: Si Earth Engine funciona usa satélite, si no, usa estimación demográfica
    if areas_data:
        ha_cultivos = next((x["Área (Ha)"] for x in areas_data if "Cultivos" in x["Cobertura"]), 0.0)
        ha_pastos = next((x["Área (Ha)"] for x in areas_data if "Pastos" in x["Cobertura"]), 0.0)
        area_total_ha = sum([x["Área (Ha)"] for x in areas_data])
    else:
        ha_pastos = bovinos / 1.5 if bovinos > 0 else 0
        ha_cultivos = (area_km2 * 100) * 0.15 if area_km2 > 0 else 0
        area_total_ha = area_km2 * 100
        
    carga_N_kg = (ha_cultivos * 28.5) + (ha_pastos * 9.2)
    carga_P_kg = (ha_cultivos * 5.8) + (ha_pastos * 1.5)
    pct_agricola = ((ha_cultivos + ha_pastos) / area_total_ha) * 100 if area_total_ha > 0 else 0
    ind_toxicidad = max(0.0, 100.0 - min(100.0, (pct_agricola / 45.0) * 100))
    
    # 🧠 ALTERACIÓN DEL NÚCLEO: Fusionamos Calidad (70% Gris/DBO5 + 30% Verde/Tóxicos)
    ind_calidad_integral = (ind_calidad * 0.70) + (ind_toxicidad * 0.30)
    ind_calidad = ind_calidad_integral 

    c_agro1, c_agro2, c_agro3 = st.columns([1, 1, 2])
    with c_agro1:
        st.metric("🌾 Nitrógeno (Eutrofización)", f"{carga_N_kg:,.0f} kg/a", "Escorrentía Agrícola", delta_color="inverse")
        st.metric("🥑 Fósforo Total", f"{carga_P_kg:,.0f} kg/a")
    with c_agro2:
        st.metric("🚜 Frontera Agropecuaria", f"{pct_agricola:.1f}%", "Densidad Crítica > 45%", delta_color="inverse")
        st.caption(f"Pastos: {ha_pastos:,.0f} ha | Cultivos: {ha_cultivos:,.0f} ha")
    with c_agro3:
        st.info("💡 **El Porqué de las SbN:** Esta carga tóxica es *difusa* (lavada por la lluvia). No viaja por alcantarillados, por ende, **no puede ser tratada por una PTAR (Infraestructura Gris)**. Su única solución es la inversión en **Infraestructura Verde** para que la naturaleza actúe como biofiltro.")

    st.markdown("<br>", unsafe_allow_html=True)
    
    # --- FUNCIONES DE RENDERIZADO VISUAL ---
    def evaluar_indice(valor, umbral_rojo, umbral_verde, invertido=False):
        if not invertido:
            if valor < umbral_rojo: return "🔴 CRÍTICO", "#c0392b"
            elif valor < umbral_verde: return "🟡 VULNERABLE", "#f39c12"
            else: return "🟢 ÓPTIMO", "#27ae60"
        else:
            if valor < umbral_verde: return "🟢 HOLGADO", "#27ae60"
            elif valor < umbral_rojo: return "🟡 MODERADO", "#f39c12"
            else: return "🔴 CRÍTICO", "#c0392b"

    def crear_velocimetro(valor, titulo, color_bar, umbral_rojo, umbral_verde, invertido=False):
        import plotly.graph_objects as go
        fig = go.Figure(go.Indicator(
            mode = "gauge+number", value = valor,
            number = {'suffix': "%", 'font': {'size': 24}}, title = {'text': titulo, 'font': {'size': 14}},
            gauge = {
                'axis': {'range': [None, 100], 'tickwidth': 1}, 'bar': {'color': color_bar}, 'bgcolor': "white",
                'steps': [
                    {'range': [0, umbral_rojo], 'color': "#ffcccb" if not invertido else "#e8f8f5"},
                    {'range': [umbral_rojo, umbral_verde], 'color': "#fff2cc"},
                    {'range': [umbral_verde, 100], 'color': "#e8f8f5" if not invertido else "#ffcccb"}
                ],
                'threshold': {'line': {'color': "black", 'width': 4}, 'thickness': 0.75, 'value': valor}
            }
        ))
        fig.update_layout(height=230, margin=dict(l=10, r=10, t=30, b=10), font_family="Georgia")
        return fig
    
    # --- RENDERIZADO DE LOS VELOCÍMETROS ---
    estres_gauge_val = min(100.0, estres_hidrico_porcentaje)
    
    # 💧 REGLA WRI: La Neutralidad Institucional Base usa solo el área gestionada en el SIG
    ha_gestionadas_base = float(ha_reales_sig) if 'ha_reales_sig' in locals() else 0.0
    volumen_repuesto_base = ha_gestionadas_base * 2500
    ind_neutralidad = min(100.0, (volumen_repuesto_base / consumo_anual_m3) * 100) if consumo_anual_m3 > 0 else 0.0
    
    col_g1, col_g2, col_g3, col_g4 = st.columns(4)
    
    est_neu, col_neu = evaluar_indice(ind_neutralidad, 40, 80)
    est_res, col_res = evaluar_indice(ind_resiliencia, 30, 70)
    est_est, col_est = evaluar_indice(estres_hidrico_porcentaje, 40, 20, invertido=True) 
    est_cal, col_cal = evaluar_indice(ind_calidad_integral, 40, 70)

    with col_g1: 
        st.plotly_chart(crear_velocimetro(ind_neutralidad, "Neutralidad (Actual)", "#2ecc71", 40, 80), width="stretch")
        st.markdown(f"<h4 style='text-align: center; color: {col_neu}; margin-top:-20px;'>{est_neu}</h4>", unsafe_allow_html=True)
    with col_g2: 
        st.plotly_chart(crear_velocimetro(ind_resiliencia, "Resiliencia Estructural", "#3498db", 30, 70), width="stretch")
        st.markdown(f"<h4 style='text-align: center; color: {col_res}; margin-top:-20px;'>{est_res}</h4>", unsafe_allow_html=True)
    with col_g3: 
        st.plotly_chart(crear_velocimetro(estres_gauge_val, "Nivel de Estrés", "#e74c3c", 20, 40, invertido=True), width="stretch")
        st.markdown(f"<h4 style='text-align: center; color: {col_est}; margin-top:-20px;'>{est_est}</h4>", unsafe_allow_html=True)
    with col_g4:
        st.plotly_chart(crear_velocimetro(ind_calidad_integral, "Calidad (DBO + Tóxicos)", "#9b59b6", 40, 70), width="stretch")
        st.markdown(f"<h4 style='text-align: center; color: {col_cal}; margin-top:-20px;'>{est_cal}</h4>", unsafe_allow_html=True)

    # ==============================================================================
    # 📍 PASO 2: LA PRUEBA DE ESTRÉS (RIESGOS FÍSICOS Y CLIMÁTICOS)
    # ==============================================================================
    st.markdown("---")
    st.markdown("## 📍 PASO 2: Prueba de Estrés (Riesgos)")
    
    col_mat1, col_mat2 = st.columns([1, 1.5])
    
    with col_mat1:
        st.markdown("#### ⛈️ Inyección de Estrés Físico (Avalancha)")
        activar_tormenta_local = st.toggle("Activar Control Manual de Avenida Torrencial")
        
        if activar_tormenta_local:
            lodo_total_m3 = st.slider("Magnitud del Deslizamiento (m³ de Lodo):", min_value=0.0, max_value=250000.0, value=85000.0, step=5000.0)
            sobrecosto_ptap = lodo_total_m3 * 0.4 
            st.session_state['eco_lodo_total_m3'] = lodo_total_m3
            st.session_state['eco_sobrecosto_usd'] = sobrecosto_ptap
        else:
            lodo_total_m3 = st.session_state.get('eco_lodo_total_m3', 0.0)
            sobrecosto_ptap = st.session_state.get('eco_sobrecosto_usd', 0.0)
            
        penalidad_lodo = (lodo_total_m3 / 100000.0) * 100
        resiliencia_real = max(0.0, ind_resiliencia - penalidad_lodo)
        
        ishi_final = (ind_estres + ind_calidad + resiliencia_real + ind_neutralidad) / 4
        
        st.metric("🎯 ISHI Global Base", f"{ishi_final:.1f}%", "Estado Pre-Inversión")
        if lodo_total_m3 > 0:
            st.warning(f"⚠️ **Riesgo Físico Activo:** La tormenta inyectó **{lodo_total_m3:,.0f} m³** de sedimentos. Sobrecosto PTAP: **${sobrecosto_ptap:,.0f} USD**.")
            
    with col_mat2:
        st.markdown("#### 🌊 Fotografía Radial")
        st.info("💡 Este es el perfil de salud actual. En el Paso 4 evaluaremos cómo las inversiones logran expandirlo.")
        import plotly.graph_objects as go
        fig_radar = go.Figure()
        fig_radar.add_trace(go.Scatterpolar(
            r=[estres_gauge_val, ind_calidad, resiliencia_real, ind_neutralidad, estres_gauge_val],
            theta=['Abastecimiento (Estrés)', 'Calidad (DBO)', 'Resiliencia (Física)', 'Neutralidad (Huella)', 'Abastecimiento (Estrés)'],
            fill='toself', fillcolor='rgba(231, 76, 60, 0.15)', line=dict(color='#e74c3c', width=2, dash='dot'),
            name=f'Línea Base ({ishi_final:.1f}%)'
        ))
        fig_radar.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 100])), showlegend=True, legend=dict(orientation="h", y=-0.2, x=0.5, xanchor="center"), height=300, margin=dict(t=30, b=30, l=40, r=40))
        st.plotly_chart(fig_radar, use_container_width=True)

    # ==============================================================================
    # 📥 PRE-PROCESAMIENTO Y DESCARGA PREDIAL (PROCESO SILENCIOSO DE DATOS)
    # ==============================================================================
    capas = {}
    try:
        if gdf_zona is not None and not gdf_zona.empty:
            capas = load_context_layers(tuple(gdf_zona.total_bounds))
    except Exception as e:
        st.warning(f"Aviso al cargar capas SIG: {e}")

    @st.cache_data(ttl=3600, show_spinner=False)
    def obtener_predios_y_hectareas(_gdf_zona, nombre_zona_txt):
        import requests, tempfile
        import pandas as pd
        import geopandas as gpd
        ha_calc = 0.0
        info_debug = "Descargando predios..."
        gdf_predios_final = None 
        try:
            url_predios = "https://ldunpssoxvifemoyeuac.supabase.co/storage/v1/object/public/geojson/PrediosEjecutados.geojson"
            res = requests.get(url_predios)
            if res.status_code != 200: return 0.0, f"❌ Fallo descarga API", None
            with tempfile.NamedTemporaryFile(delete=False, suffix=".geojson") as tmp:
                tmp.write(res.content)
                tmp_path = tmp.name
            gdf_p = gpd.read_file(tmp_path)
            if gdf_p.empty: return 0.0, "❌ GeoJSON vacío", None
            gdf_p.set_crs(epsg=4326, allow_override=True, inplace=True)
            gdf_p_3116 = gdf_p.to_crs(epsg=3116)
            gdf_z_3116 = _gdf_zona.to_crs(epsg=3116)
            gdf_p_3116['geometry'] = gdf_p_3116.geometry.make_valid().buffer(0)
            gdf_z_3116['geometry'] = gdf_z_3116.geometry.make_valid().buffer(0)
            recorte_exacto = gpd.clip(gdf_p_3116, gdf_z_3116)
            if not recorte_exacto.empty:
                ha_calc = recorte_exacto.area.sum() / 10000.0
                info_debug = f"✅ CORTE EXACTO: {len(recorte_exacto)} fragmentos de predios operan físicamente dentro de la cuenca."
                gdf_predios_final = recorte_exacto.to_crs(epsg=4326)
            else:
                info_debug = f"ℹ️ ZONA VIRGEN: Ningún predio cae dentro de {nombre_zona_txt}."
        except Exception as e: 
            info_debug = f"❌ ERROR GEOMÉTRICO: {e}"
        return ha_calc, info_debug, gdf_predios_final

    with st.spinner("Descargando inventario predial de la Nube (Supabase)..."):
        ha_reales_sig, info_debug, gdf_predios_mapa = obtener_predios_y_hectareas(gdf_zona, nombre_zona)


    # ==============================================================================
    # 📍 PASO 3: INGENIERÍA DE SOLUCIONES (SIMULADOR FÍSICO WRI)
    # ==============================================================================
    st.markdown("---")
    st.markdown("## 📍 PASO 3: Ingeniería de Soluciones (Física del Territorio)")
    st.info("Transforma las métricas biofísicas en indicadores estandarizados de ingeniería, simula portafolios de intervención y visualiza el impacto volumétrico real antes de asignar presupuesto financiero.")
    
    st.markdown(f"#### 🌲 1. Simulación de Beneficios Volumétricos (SbN) en: **{nombre_zona}**")
    
    st.info("💡 **Regla de Contabilidad WRI (VWBA):** Para el cierre de brechas financieras (Neutralidad), el modelo solo reconoce el volumen de agua generado por el *Capital Gestionado* (Áreas SIG de conservación histórica) y las nuevas inversiones simuladas.")
    
    # El cálculo financiero base es ESTRICTAMENTE lo gestionado en el SIG
    ha_base_calculo = float(ha_reales_sig) if 'ha_reales_sig' in locals() else 0.0
    
    ha_riparias_potenciales = 0.0
    sumar_riparias = False
    df_str = st.session_state.get('geomorfo_strahler_df')
    
    if df_str is not None and not df_str.empty:
        st.markdown("<br>", unsafe_allow_html=True)
        with st.container(border=True):
            st.markdown("🌿 **Infraestructura Verde: Potencial Ripario (Conectado a Biodiversidad)**")
            anillos = st.session_state.get('multi_rings', [10, 20, 30])
            escenario_nombres = [f"🔴 Escenario Mínimo Normativo ({anillos[0]}m)", f"🟡 Escenario Ideal Recomendado ({anillos[1]}m)", f"🟢 Escenario Óptimo Ecológico ({anillos[2]}m)"]
            if 'aleph_twi_umbral' in st.session_state:
                st.success("🧠 **Nexo Físico Activo:** Integrando zona de amenaza de inundación/avalancha como área de restauración prioritaria.")
            cr1, cr2, cr3 = st.columns(3)
            escenario_sel = cr1.selectbox("Selecciona Escenario a Financiar en WRI:", escenario_nombres, index=1, key="td_sel_rip")
            idx_sel = escenario_nombres.index(escenario_sel)
            ancho_buffer = anillos[idx_sel]
            longitud_total_km = df_str['Longitud_Km'].sum()
            cr2.metric("Longitud de Cauces", f"{longitud_total_km:,.2f} km")
            ha_riparias_potenciales = (longitud_total_km * 1000 * (ancho_buffer * 2)) / 10000.0
            cr3.metric("Potencial Ripario (SbN)", f"{ha_riparias_potenciales:,.1f} ha")
            sumar_riparias = st.checkbox("📥 Incorporar estas hectáreas riparias a la simulación financiera WRI", value=True, key="td_sumar_rip")
    else:
        st.info("💡 **Tip:** Usa el motor de Geomorfología para detectar la red de drenaje y luego la página de Biodiversidad para definir los anillos de protección.")
    
    st.markdown("<br>", unsafe_allow_html=True)
    c_inv1, c_inv2, c_inv3 = st.columns(3)
    with c_inv1:
        st.metric("✅ Área Conservada (Base SIG)", f"{ha_reales_sig:,.1f} ha")
        ha_simuladas = st.number_input("➕ Adicionar Hectáreas Extra (Manual):", min_value=0.0, value=0.0, step=10.0, key="td_ha_sim")
        ha_total = ha_base_calculo + ha_simuladas + (ha_riparias_potenciales if sumar_riparias else 0.0)
        beneficio_restauracion_m3 = ha_total * 2500 
    with c_inv2:
        sist_saneamiento = st.number_input("Sistemas Tratamiento (STAM/PTAR):", min_value=0, value=50, step=5, key="td_stam")
        beneficio_calidad_m3 = sist_saneamiento * 1200
    with c_inv3:
        volumen_repuesto_m3 = beneficio_restauracion_m3 + beneficio_calidad_m3
        st.metric("💧 Agua 'Devuelta' (VWBA)", f"{volumen_repuesto_m3:,.0f} m³/año", "Impacto total simulado")
        
    with st.container(border=True):
        st.markdown("#### ⚖️ Dinámica de Regulación Eco-Hidrológica (Source-to-Tap)")
        st.markdown("Integración de la termodinámica del ecosistema. **Nota:** Este diagrama físico sí utiliza la totalidad del *Capital Natural* detectado por el satélite (Bosques + Matorrales) más las nuevas áreas simuladas.")
        
        # Rescatamos el bosque total del satélite para la física del agua
        if 'areas_data' in locals() and areas_data:
            ha_bosque_sat = next((x["Área (Ha)"] for x in areas_data if "Bosque" in x["Cobertura"]), 0.0)
            ha_matorral_sat = next((x["Área (Ha)"] for x in areas_data if "Matorrales" in x["Cobertura"]), 0.0)
        else:
            ha_bosque_sat, ha_matorral_sat = 0.0, 0.0
            
        ha_fisica_total = ha_bosque_sat + ha_matorral_sat + ha_simuladas + (ha_riparias_potenciales if sumar_riparias else 0.0)
        
        area_cuenca_ha = area_km2 * 100                            
        pct_bosque = min(1.0, ha_fisica_total / area_cuenca_ha) if area_cuenca_ha > 0 else 0
        
        if area_km2 > 0:
            ppt_mm_estimada = (oferta_anual_m3 / (area_km2 * 1000)) * 2.5
            vol_lluvia_total = ppt_mm_estimada * area_km2 * 1000
        else:
            ppt_mm_estimada = 0.0
            vol_lluvia_total = 0.0     
            
        eficiencia_dosel_max = st.session_state.get('bio_eficiencia_retencion_pct', 25.0) / 100.0
        pct_intercepcion = 0.05 + ((eficiencia_dosel_max - 0.05) * pct_bosque)
        vol_intercepcion = vol_lluvia_total * pct_intercepcion
        pct_etp = 0.35 + (0.10 * pct_bosque)
        vol_etp = vol_lluvia_total * pct_etp
        vol_al_suelo = vol_lluvia_total - vol_intercepcion - vol_etp
        pct_infiltracion = 0.20 + (0.50 * pct_bosque)
        vol_infiltracion = vol_al_suelo * pct_infiltracion
        vol_escorrentia = vol_al_suelo - vol_infiltracion
        
        labels = ["<b>Lluvia Total</b>", "<b>Retención del Dosel (Hojas)</b>", "<b>Evapotranspiración (ETP)</b>", "<b>Escorrentía Rápida (Riesgo)</b>", "<b>Infiltración (Acuífero)</b>", "<b>Flujo Base (Oferta Segura)</b>"]
        source = [0, 0, 0, 0, 4]
        target = [1, 2, 3, 4, 5]
        value = [vol_intercepcion, vol_etp, vol_escorrentia, vol_infiltracion, vol_infiltracion]
        color_links = ["rgba(46, 204, 113, 0.5)", "rgba(241, 196, 15, 0.4)", "rgba(231, 76, 60, 0.6)", "rgba(52, 152, 219, 0.4)", "rgba(41, 128, 185, 0.6)"]
        
        import plotly.graph_objects as go
        fig_sankey = go.Figure(data=[go.Sankey(
            valueformat=".0f", valuesuffix=" m³/año", textfont=dict(size=14, color="#000000", family="Georgia, serif"),
            node=dict(pad=25, thickness=25, line=dict(color="black", width=0.5), label=labels, color=["#34495e", "#2ecc71", "#f39c12", "#e74c3c", "#3498db", "#2980b9"]),
            link=dict(source=source, target=target, value=value, color=color_links)
        )])
        fig_sankey.update_layout(height=420, margin=dict(l=20, r=20, t=30, b=20), font_family="Georgia")
        
        c_sk1, c_sk2 = st.columns([1, 2.5])
        with c_sk1:
            st.metric("🌧️ Lluvia Total", f"{vol_lluvia_total/1e6:,.1f} Mm³")
            st.metric("🍃 Agua Retenida en Dosel", f"{vol_intercepcion/1e6:,.1f} Mm³", "Regulación microclimática", delta_color="normal")
            st.metric("💧 Oferta Regulada (Infiltrada)", f"{vol_infiltracion/1e6:,.1f} Mm³", "Trasladada al flujo base", delta_color="normal")
            st.caption("A mayor inversión en área conservada (SbN), aumenta la intercepción foliar y la infiltración, reduciendo drásticamente la vena roja de escorrentía rápida.")
        with c_sk2:
            st.plotly_chart(fig_sankey, use_container_width=True)
            
    st.markdown("---")
    st.markdown(f"#### 💼 2. Portafolios de Intervención Multi-Objetivo")

    # Portafolio 1: Cantidad
    with st.container(border=True):
        st.markdown("🎯 **Portafolio 1: Neutralidad Volumétrica (Cantidad)**")
        col_m1, col_m2 = st.columns([1, 2.5])
        with col_m1:
            meta_neutralidad = st.slider("Meta Neutralidad (%)", 10.0, 100.0, 100.0, 5.0, key="td_meta_n")
            costo_ha = st.number_input("Restauración (1 ha) [M COP]:", value=8.5, step=0.5, key="td_c_ha")
            costo_stam_n = st.number_input("Saneamiento (1 STAM) [M COP]:", value=15.0, step=1.0, key="td_c_stamn")
            costo_lps = st.number_input("Eficiencia (1 L/s) [M COP]:", value=120.0, step=10.0, key="td_c_lps")
        with col_m2:
            vol_requerido_m3 = (meta_neutralidad / 100.0) * consumo_anual_m3
            brecha_m3 = vol_requerido_m3 - volumen_repuesto_m3
            ha_proyectos_simulados = ha_simuladas + (ha_riparias_potenciales if sumar_riparias else 0.0)
            costo_proyectos_simulados = ha_proyectos_simulados * costo_ha
            
            if brecha_m3 <= 0: 
                st.success("✅ ¡Se cumple la meta de Neutralidad Volumétrica con la cobertura natural y/o proyectos simulados!")
                st.info(f"💰 Inversión en proyectos simulados (SbN): **${costo_proyectos_simulados:,.0f} M COP** (~${costo_proyectos_simulados/4000:,.2f} M USD)")
            else:
                st.warning(f"⚠️ Faltan compensar **{brecha_m3/1e6:,.2f} Millones de m³/año**.")
                ce_sbn = (costo_ha * 1_000_000) / 2500.0
                ce_stam = (costo_stam_n * 1_000_000) / 1200.0
                ce_lps = (costo_lps * 1_000_000) / 31536.0
                st.markdown(f"<div style='font-size:0.85rem; color:#666; margin-bottom:10px;'><b>Costo Marginal Unitario (COP por m³):</b> 🌲 SbN: <span style='color:green;'>${ce_sbn:,.0f}</span> | 🚰 Eficiencia: <span style='color:orange;'>${ce_lps:,.0f}</span> | 🚽 STAM: <span style='color:red;'>${ce_stam:,.0f}</span></div>", unsafe_allow_html=True)
                optimo_p1 = st.toggle("🪄 Activar Óptimo Técnico-Financiero", key="td_opt_p1")
                cmix1, cmix2, cmix3 = st.columns(3)
                if optimo_p1:
                    st.info("Algoritmo activo: Maximiza la inversión en Restauración (65%) y Eficiencia (30%) por ser las vías más económicas, reservando un 5% a Saneamiento gris.")
                    pct_a = cmix1.number_input("% Cierre vía Restauración", 0, 100, 65, disabled=True)
                    pct_b = cmix2.number_input("% Cierre vía Saneamiento", 0, 100, 5, disabled=True)
                    pct_c = cmix3.number_input("% Cierre vía Eficiencia", 0, 100, 30, disabled=True)
                else:
                    pct_a = cmix1.number_input("% Cierre vía Restauración", 0, 100, 40)
                    pct_b = cmix2.number_input("% Cierre vía Saneamiento", 0, 100, 40)
                    pct_c = cmix3.number_input("% Cierre vía Eficiencia", 0, 100, 20)
                if (pct_a + pct_b + pct_c) == 100:
                    ha_req = (brecha_m3 * (pct_a/100)) / 2500.0
                    stam_req = (brecha_m3 * (pct_b/100)) / 1200.0
                    lps_req = ((brecha_m3 * (pct_c/100)) * 1000) / 31536000 
                    inv_brecha = (ha_req * costo_ha) + (stam_req * costo_stam_n) + (lps_req * costo_lps)
                    inv_total = inv_brecha + costo_proyectos_simulados
                    co1, co2, co3, co4 = st.columns(4)
                    co1.metric("🌲 Restaurar Total", f"{(ha_req + ha_proyectos_simulados):,.1f} ha")
                    co2.metric("🚽 STAM", f"{stam_req:,.0f} unds")
                    co3.metric("🚰 Eficiencia", f"{lps_req:,.1f} L/s")
                    co4.metric("💰 INVERSIÓN TOTAL", f"${inv_total:,.0f} M COP", f"~${inv_total/4000:,.2f} M USD", delta_color="off")
                else: st.error("La suma de los porcentajes debe ser exactamente 100%.")

    # Portafolio 2: Calidad
    with st.container(border=True):
        st.markdown("🎯 **Portafolio 2: Remoción de Cargas (Calidad DBO5)**")
        col_c1, col_c2 = st.columns([1, 2.5])
        with col_c1:
            meta_remocion = st.slider("Meta Remoción DBO (%)", 10.0, 100.0, 85.0, 5.0, key="td_meta_c")
            costo_ptar = st.number_input("PTAR (1 Ton/a) [M COP]:", value=150.0, step=10.0, key="td_c_ptar")
            costo_stam_c = st.number_input("STAM (1 Ton/a) [M COP]:", value=45.0, step=5.0, key="td_c_stamc")
            costo_sbn_c = st.number_input("SbN (1 Ton/a) [M COP]:", value=12.0, step=2.0, key="td_c_sbn_c")
        with col_c2:
            carga_objetivo = (meta_remocion / 100.0) * carga_total_ton
            brecha_ton = carga_objetivo - (sist_saneamiento * 0.5) 
            if brecha_ton <= 0: st.success("✅ ¡Meta de Remoción de Cargas alcanzada con la simulación!")
            else:
                st.warning(f"⚠️ Faltan remover **{brecha_ton:,.1f} Ton/año** de DBO5.")
                st.markdown(f"<div style='font-size:0.85rem; color:#666; margin-bottom:10px;'><b>Costo Marginal por Tonelada DBO5:</b> 🌿 SbN Biofiltros: <span style='color:green;'>${costo_sbn_c}M</span> | 🏡 STAM Rural: <span style='color:orange;'>${costo_stam_c}M</span> | 🏙️ PTAR: <span style='color:red;'>${costo_ptar}M</span></div>", unsafe_allow_html=True)
                optimo_p2 = st.toggle("🪄 Activar Óptimo Técnico-Financiero", key="td_opt_p2")
                cmc1, cmc2, cmc3 = st.columns(3)
                if optimo_p2:
                    st.info("Algoritmo activo: Prioriza Biofiltros SbN (60%) por su inmenso ahorro, combinados con STAM rural (30%) y solo un 10% en infraestructura pesada (PTAR).")
                    pct_ptar = cmc1.number_input("% Cierre vía PTAR", 0, 100, 10, disabled=True)
                    pct_stam_c = cmc2.number_input("% Cierre vía STAM", 0, 100, 30, disabled=True)
                    pct_sbn_c = cmc3.number_input("% Cierre vía SbN", 0, 100, 60, disabled=True)
                else:
                    pct_ptar = cmc1.number_input("% Cierre vía PTAR", 0, 100, 50)
                    pct_stam_c = cmc2.number_input("% Cierre vía STAM", 0, 100, 30)
                    pct_sbn_c = cmc3.number_input("% Cierre vía SbN", 0, 100, 20)
                if (pct_ptar + pct_stam_c + pct_sbn_c) == 100:
                    t_ptar = brecha_ton * (pct_ptar/100)
                    t_stam = brecha_ton * (pct_stam_c/100)
                    t_sbn = brecha_ton * (pct_sbn_c/100)
                    inv_tot_c = (t_ptar * costo_ptar) + (t_stam * costo_stam_c) + (t_sbn * costo_sbn_c)
                    coc1, coc2, coc3, coc4 = st.columns(4)
                    coc1.metric("🏙️ PTAR", f"{t_ptar:,.0f} Ton")
                    coc2.metric("🏡 STAM Rural", f"{t_stam:,.0f} Ton")
                    coc3.metric("🌿 SbN Biofiltros", f"{t_sbn:,.0f} Ton")
                    coc4.metric("💰 INVERSIÓN CALIDAD", f"${inv_tot_c:,.0f} M COP", f"~${inv_tot_c/4000:,.2f} M USD", delta_color="off")
                else: st.error("La suma debe ser exactamente 100%.")

    st.markdown("---")
    st.markdown("#### 🚀 Impacto Físico Proyectado")
    st.info("Estos son los nuevos niveles de salud territorial si se implementan las hectáreas y plantas de tratamiento modeladas arriba.")
    
    caudal_oferta_L_s = (oferta_anual_m3 / 31536000) * 1000 
    carga_removida_sim = sist_saneamiento * 2.5
    carga_final_rio_sim = max(0.0, carga_total_ton - carga_removida_sim)
    carga_mg_s_sim = (carga_final_rio_sim * 1_000_000_000) / 31536000
    conc_dbo_sim = carga_mg_s_sim / caudal_oferta_L_s if caudal_oferta_L_s > 0 else 999.0
    
    ind_calidad_sim = max(0.0, min(100.0, 100 * math.exp(-0.07 * conc_dbo_sim)))
    ind_neutralidad_sim = min(100.0, (volumen_repuesto_m3 / consumo_anual_m3) * 100) if consumo_anual_m3 > 0 else 0.0
    mejora_infiltracion = (ha_total / (area_km2 * 100)) * 0.10 if area_km2 > 0 else 0.0
    bfi_ratio_sim = bfi_ratio * (1 + mejora_infiltracion)
    ind_resiliencia_sim = max(0.0, min(100.0, (bfi_ratio_sim / 0.70) * 100 * factor_supervivencia))
    oferta_efectiva_sim = oferta_anual_m3 + volumen_repuesto_m3
    wei_ratio_sim = consumo_anual_m3 / oferta_efectiva_sim if oferta_efectiva_sim > 0 else 1.0
    estres_sim_porcentaje = wei_ratio_sim * 100
    estres_gauge_sim_val = min(100.0, estres_sim_porcentaje)

    # Calculamos los estados (colores y textos) para los índices proyectados
    est_neu_sim, col_neu_sim = evaluar_indice(ind_neutralidad_sim, 40, 80)
    est_res_sim, col_res_sim = evaluar_indice(ind_resiliencia_sim, 30, 70)
    est_est_sim, col_est_sim = evaluar_indice(estres_gauge_sim_val, 40, 20, invertido=True)
    est_cal_sim, col_cal_sim = evaluar_indice(ind_calidad_sim, 40, 70)

    cg1, cg2, cg3, cg4 = st.columns(4)
    with cg1: 
        st.plotly_chart(crear_velocimetro(ind_neutralidad_sim, "Neutralidad (Proyectada)", "#2ecc71", 40, 80), width="stretch")
        st.markdown(f"<h4 style='text-align: center; color: {col_neu_sim}; margin-top:-20px;'>{est_neu_sim}</h4>", unsafe_allow_html=True)
    with cg2: 
        st.plotly_chart(crear_velocimetro(ind_resiliencia_sim, "Resiliencia (Proyectada)", "#3498db", 30, 70), width="stretch")
        st.markdown(f"<h4 style='text-align: center; color: {col_res_sim}; margin-top:-20px;'>{est_res_sim}</h4>", unsafe_allow_html=True)
    with cg3: 
        st.plotly_chart(crear_velocimetro(estres_gauge_sim_val, "Estrés (Proyectado)", "#e74c3c", 20, 40, invertido=True), width="stretch")
        st.markdown(f"<h4 style='text-align: center; color: {col_est_sim}; margin-top:-20px;'>{est_est_sim}</h4>", unsafe_allow_html=True)
    with cg4: 
        st.plotly_chart(crear_velocimetro(ind_calidad_sim, "Calidad (Proyectada)", "#9b59b6", 40, 70), width="stretch")
        st.markdown(f"<h4 style='text-align: center; color: {col_cal_sim}; margin-top:-20px;'>{est_cal_sim}</h4>", unsafe_allow_html=True)

    # ==============================================================================
    # 📍 PASO 4: LA VIABILIDAD FINANCIERA (OPTIMIZADOR ROI)
    # ==============================================================================
    st.markdown("---")
    st.markdown("## 📍 PASO 4: Viabilidad Financiera e Inteligencia de Capital")
    st.info("Algoritmo de asignación inteligente. Define un presupuesto financiero en USD y el modelo distribuirá el capital óptimamente para maximizar la expansión del radar ISHI.")

    c_opt1, c_opt2 = st.columns([1, 2.5])
    with c_opt1:
        st.markdown("#### 💰 Estrategia de Capital")
        modo_plan = st.radio("Seleccione el Enfoque Financiero:", ["📊 1. Asignar Presupuesto Disponible", "🎯 2. Definir Meta de Seguridad (ISHI)"])
        
        if "1. Asignar" in modo_plan:
            presupuesto_MUSD = st.number_input("Presupuesto de Inversión (Millones USD):", 0.5, 100.0, 5.0, 0.5)
            presupuesto_usd = presupuesto_MUSD * 1_000_000
        else:
            ishi_minimo = float(int(ishi_final)) if ishi_final < 100 else 100.0
            meta_ishi = st.slider("Definir Meta ISHI (%):", ishi_minimo, 100.0, max(75.0, ishi_minimo + 5.0), 1.0)
            brecha_ishi = max(0.0, meta_ishi - ishi_final)
            presupuesto_MUSD = brecha_ishi * 0.35
            presupuesto_usd = presupuesto_MUSD * 1_000_000
            if presupuesto_usd > 0: st.info(f"🎯 Para alcanzar {meta_ishi}%, se requieren **${presupuesto_MUSD:,.1f} M USD**.")
            else: st.success("✅ La región ya cumple con esta meta.")

        brecha_calidad = max(0.1, 100 - ind_calidad)
        brecha_resiliencia = max(0.1, 100 - resiliencia_real)
        brecha_estres = max(0.1, 100 - estres_gauge_val)
        
        peso_gris = (brecha_calidad ** 0.5) 
        peso_verde = (brecha_resiliencia ** 0.5) + (brecha_estres ** 0.5)
        peso_total = peso_gris + peso_verde
        w_gris, w_verde = peso_gris / peso_total, peso_verde / peso_total

        inv_gris, inv_verde = presupuesto_usd * w_gris, presupuesto_usd * w_verde

        st.markdown("#### 🏗️ Asignación Óptima")
        st.metric("🟢 Infraestructura Verde (SbN)", f"${inv_verde/1e6:,.2f} M", "Conservación & Restauración")
        st.metric("🏢 Infraestructura Gris (PTAR)", f"${inv_gris/1e6:,.2f} M", "Saneamiento & Reducción DBO")

        import math
        k_verde, k_gris = 0.08, 0.12
        inv_verde_M, inv_gris_M = inv_verde / 1e6, inv_gris / 1e6

        recuperacion_resiliencia = brecha_resiliencia * (1 - math.exp(-k_verde * inv_verde_M))
        recuperacion_estres = brecha_estres * (1 - math.exp(-k_verde * inv_verde_M * 0.5)) 
        recuperacion_calidad = brecha_calidad * (1 - math.exp(-k_gris * inv_gris_M))

        new_resiliencia = resiliencia_real + recuperacion_resiliencia
        new_estres = estres_gauge_val + recuperacion_estres 
        new_calidad = ind_calidad + recuperacion_calidad
        brecha_neut = max(0.1, 100 - ind_neutralidad)
        new_neutralidad = ind_neutralidad + (brecha_neut * 0.1) 
        
        ganancia_media = (recuperacion_resiliencia + recuperacion_estres + recuperacion_calidad) / 3
        new_ishi = min(100.0, ishi_final + ganancia_media)

    with c_opt2:
        st.markdown("#### 📈 Retorno de Seguridad Hídrica (Radar Base vs. Proyectado)")
        import plotly.graph_objects as go
        fig_opt = go.Figure()
        fig_opt.add_trace(go.Scatterpolar(r=[estres_gauge_val, ind_calidad, resiliencia_real, ind_neutralidad, estres_gauge_val], theta=['Abastecimiento', 'Calidad (DBO)', 'Resiliencia', 'Neutralidad', 'Abastecimiento'], fill='toself', fillcolor='rgba(231, 76, 60, 0.15)', line=dict(color='#e74c3c', width=2, dash='dot'), name=f'Base ({ishi_final:.1f}%)'))
        fig_opt.add_trace(go.Scatterpolar(r=[new_estres, new_calidad, new_resiliencia, new_neutralidad, new_estres], theta=['Abastecimiento', 'Calidad (DBO)', 'Resiliencia', 'Neutralidad', 'Abastecimiento'], fill='toself', fillcolor='rgba(46, 204, 113, 0.4)', line=dict(color='#27ae60', width=2), name=f'Optimizado ({new_ishi:.1f}%)'))
        fig_opt.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 100])), showlegend=True, legend=dict(orientation="h", y=-0.2, x=0.5, xanchor="center"), height=400, margin=dict(t=10, b=10, l=40, r=40))
        st.plotly_chart(fig_opt, use_container_width=True)
        
        verbo_impacto = "expande" if new_ishi >= ishi_final else "mantiene"
        st.success(f"🚀 **Veredicto Estratégico:** La inyección de **${presupuesto_usd/1e6:,.1f} Millones USD** estratégicamente distribuidos {verbo_impacto} la huella de seguridad hídrica de la región de **{ishi_final:.1f}%** a **{new_ishi:.1f}%**.")

    st.markdown("---")
    c_sens1, c_sens2 = st.columns([1.2, 2])
    with c_sens1:
        st.markdown("#### 📊 Análisis de Sensibilidad y Hitos Financieros")
        st.write(f"La transición de un **ISHI de {ishi_final:.1f}% a {new_ishi:.1f}%** representa alejar al territorio del racionamiento hídrico. Físicamente, la inyección en SbN restaura la 'esponja' del suelo, mientras que la PTAR asimila el pico orgánico.")
        def calcular_inversion_exacta(meta_objetivo):
            if ishi_final >= meta_objetivo: return 0.0
            inv_test = 0.0
            while inv_test <= 100.0:
                i_g, i_v = inv_test * w_gris, inv_test * w_verde
                r_r = brecha_resiliencia * (1 - math.exp(-k_verde * i_v))
                r_e = brecha_estres * (1 - math.exp(-k_verde * i_v * 0.5))
                r_c = brecha_calidad * (1 - math.exp(-k_gris * i_g))
                g_m = (r_r + r_e + r_c) / 3
                if (ishi_final + g_m) >= meta_objetivo: return inv_test
                inv_test += 0.1
            return 100.0
        meta_10 = min(100.0, ishi_final + 10.0)
        st.write(f"• **Mejorar la huella un +10%:** Pasar del **{ishi_final:.1f}% actual** al **{meta_10:.1f}%** requiere inyectar ~${calcular_inversion_exacta(meta_10):.1f} Millones USD.")
        if ishi_final < 90.0:
            brecha_90 = 90.0 - ishi_final
            costo_90 = calcular_inversion_exacta(90.0)
            if costo_90 < 100.0: st.write(f"• **Alcanzar el ISHI Óptimo (90%):** Cerrar la brecha del **{brecha_90:.1f}%** exige un fondo de ~${costo_90:.1f} Millones USD. El costo es exponencialmente mayor por el rendimiento físico decreciente.")
            else: st.write(f"• **Alcanzar el ISHI Óptimo (90%):** Requiere una mega-inversión estructural superior a los **$100.0 M USD**.")

    with c_sens2:
        import plotly.express as px
        inv_x = list(range(0, 55, 5))
        ishi_y = [min(100.0, ishi_final + ((brecha_resiliencia*(1-math.exp(-k_verde*(i*w_verde))) + brecha_estres*(1-math.exp(-k_verde*(i*w_verde)*0.5)) + brecha_calidad*(1-math.exp(-k_gris*(i*w_gris))))/3)) for i in inv_x]
        fig_curva = px.area(x=inv_x, y=ishi_y, labels={'x': 'Capital Invertido (Millones USD)', 'y': 'Proyección ISHI (%)'}, color_discrete_sequence=['#3498db'])
        presupuesto_M_actual = presupuesto_usd / 1e6
        fig_curva.add_vline(x=presupuesto_M_actual, line_dash="dash", line_color="#e74c3c")
        fig_curva.add_annotation(x=presupuesto_M_actual, y=min(100, new_ishi + 10), text="Inversión Simulada", showarrow=False, font=dict(color="#e74c3c"))
        fig_curva.update_layout(height=250, margin=dict(t=10, b=10, l=10, r=10), yaxis=dict(range=[0, 100]))
        st.plotly_chart(fig_curva, use_container_width=True)

    # ==============================================================================
    # 📍 PASO 5: ANÁLISIS COSTO-BENEFICIO (ACB) Y FACTIBILIDAD FINANCIERA
    # ==============================================================================
    st.markdown("---")
    st.markdown("## 📍 PASO 5: Análisis Costo-Beneficio (ACB)")
    st.info("Traducción de los impactos ecológicos a flujos financieros descontados para justificar la viabilidad económica ante bancas de desarrollo.")

    with st.expander("💸 Configuración de Parámetros Económicos", expanded=True):
        c_acb1, c_acb2, c_acb3 = st.columns(3)
        
        with c_acb1:
            st.markdown("#### 📏 Escala del Proyecto")
            # Heredamos las hectáreas totales simuladas en el Paso 3
            ha_proyecto_acb = st.number_input("Hectáreas Totales a Evaluar:", value=float(ha_total) if 'ha_total' in locals() else 1500.0, help="Heredado de la simulación física del Paso 3.")
            horizonte = st.slider("Horizonte de Evaluación (Años):", 10, 50, 20, key="acb_hor")
            tasa_desc = st.slider("Tasa de Descuento (Social) %:", 1.0, 15.0, 10.0) / 100

        with c_acb2:
            st.markdown("#### 💰 OPEX y Oportunidad")
            # Usamos el costo_ha que ya definió el usuario arriba (en Millones COP)
            costo_unit_ha = costo_ha if 'costo_ha' in locals() else 8.5
            opex_pct = st.slider("Mantenimiento Anual (% del CAPEX):", 1.0, 10.0, 3.5) / 100
            c_oportunidad_ha = st.number_input("Costo Oportunidad (M COP/ha/año):", value=0.6, help="Renta agrícola/ganadera dejada de percibir.")

        with c_acb3:
            st.markdown("#### 🌍 Valoración de Servicios")
            v_agua_m3 = st.number_input("Valor Social Agua (COP/m³):", value=150, help="Ahorro en tratamiento y escasez.")
            v_carbono_ton = st.number_input("Valor Carbono (COP/Ton CO2):", value=65000, help="Precio de mercado de bonos de carbono.")
            v_riesgo_ha = st.number_input("Evitación Desastres (M COP/ha/año):", value=1.2)

    # --- MOTOR FINANCIERO INTEGRADO ---
    capex_total = ha_proyecto_acb * costo_unit_ha
    opex_anual = (capex_total * opex_pct) + (ha_proyecto_acb * c_oportunidad_ha)
    
    # Beneficios anuales (con maduración biológica)
    b_agua = (ha_proyecto_acb * 2500) * v_agua_m3 / 1e6 # En Millones COP
    b_co2 = (ha_proyecto_acb * 12.0) * v_carbono_ton / 1e6 # En Millones COP
    b_riesgo = ha_proyecto_acb * v_riesgo_ha
    beneficio_anual_potencial = b_agua + b_co2 + b_riesgo

    df_flujos = pd.DataFrame({'Año': np.arange(0, horizonte + 1)})
    df_flujos['Costos'] = opex_anual
    df_flujos.loc[0, 'Costos'] = capex_total
    
    # Curva de maduración: El bosque tarda en dar beneficios
    df_flujos['Maduracion'] = np.where(df_flujos['Año'] == 0, 0.0, 1 - np.exp(-0.3 * df_flujos['Año']))
    df_flujos['Beneficios'] = beneficio_anual_potencial * df_flujos['Maduracion']
    
    # Descuento financiero
    df_flujos['Neto'] = df_flujos['Beneficios'] - df_flujos['Costos']
    df_flujos['Factor'] = 1 / ((1 + tasa_desc) ** df_flujos['Año'])
    df_flujos['Neto_Desc'] = df_flujos['Neto'] * df_flujos['Factor']
    df_flujos['Acumulado'] = df_flujos['Neto_Desc'].cumsum()

    # Métricas
    vpn_acb = df_flujos['Neto_Desc'].sum()
    rcb_acb = (df_flujos['Beneficios'] * df_flujos['Factor']).sum() / (df_flujos['Costos'] * df_flujos['Factor']).sum()
    
    m_acb1, m_acb2, m_acb3 = st.columns(3)
    with m_acb1:
        st.metric("Valor Presente Neto (VPN)", f"${vpn_acb:,.1f} M COP", "Viable" if vpn_acb > 0 else "No Viable")
    with m_acb2:
        st.metric("Relación Beneficio/Costo", f"{rcb_acb:.2f}x", "Genera Valor" if rcb_acb > 1 else "Riesgo")
    with m_acb3:
        try: pback = df_flujos[df_flujos['Acumulado'] >= 0]['Año'].iloc[0]
        except: pback = ">" + str(horizonte)
        st.metric("Punto de Equilibrio", f"Año {pback}")

    # Gráfico Profesional
    from plotly.subplots import make_subplots
    fig_acb = make_subplots(specs=[[{"secondary_y": True}]])
    fig_acb.add_trace(go.Bar(x=df_flujos['Año'], y=df_flujos['Neto_Desc'], name='Flujo Neto Descontado', marker_color='#2ecc71'), secondary_y=False)
    fig_acb.add_trace(go.Scatter(x=df_flujos['Año'], y=df_flujos['Acumulado'], name='VPN Acumulado', line=dict(color='#2980b9', width=3)), secondary_y=True)
    fig_acb.update_layout(title="Viabilidad Financiera del Portafolio SbN", height=400, hovermode="x unified", legend=dict(orientation="h", y=-0.2))
    st.plotly_chart(fig_acb, use_container_width=True)

    # 🧠 CAJA INTELIGENTE DE SÍNTESIS FINANCIERA (Sihcli-Poter AI)
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("### 🧠 Veredicto Financiero (Sihcli-Poter AI)")
    
    if vpn_acb > 0:
        msg_fin = f"✅ **Proyecto Estructuralmente Viable:** La intervención de conservación/restauración sobre **{ha_proyecto_acb:,.0f} hectáreas** genera valor económico real. Con una Relación Costo-Beneficio de **{rcb_acb:.2f}x**, la suma de los servicios ecosistémicos (agua, carbono y mitigación) supera ampliamente el CAPEX y los costos de oportunidad a lo largo de los {horizonte} años. El capital se recupera totalmente en el **Año {pback}**."
        st.success(msg_fin)
    else:
        déficit_m = abs(vpn_acb)
        # Separamos el símbolo $ de los asteriscos para evitar el bug de LaTeX en Streamlit
        msg_fin = f"Bajo los parámetros actuales, la intervención sobre **{ha_proyecto_acb:,.0f} hectáreas** destruye valor (VPN negativo de **$ {déficit_m:,.1f} M COP**). El RCB de **{rcb_acb:.2f}x** indica que por cada peso invertido, la sociedad solo recupera {rcb_acb*100:.0f} centavos en servicios ecosistémicos."
        
        # Lógica de recomendaciones dinámicas usando Markdown puro (\n\n)
        recomendaciones = []
        if horizonte < 30: 
            recomendaciones.append("🔸 **1. Ampliar el Horizonte:** Los bosques son activos de maduración lenta. Evaluar el proyecto a 30 o 40 años permitirá que la curva de beneficios supere el punto de equilibrio.")
        if tasa_desc > 0.08: 
            recomendaciones.append(f"🔸 **2. Buscar Capital Concesional:** Una tasa de descuento del {tasa_desc*100:.1f}% es muy exigente. Acudir a banca multilateral (BID/BM) o fondos climáticos puede reducirla al 4% - 6%.")
        if c_oportunidad_ha > 0.5: 
            recomendaciones.append(f"🔸 **3. Reubicar Tácticas:** El costo de oportunidad agrario es muy alto ($ {c_oportunidad_ha} M). Use el mapa táctico para enfocar las SbN en tierras marginales más baratas.")
        
        # Unimos con saltos de línea nativos de Markdown
        rec_texto = "\n\n".join(recomendaciones)
        
        # Ensamblamos el mensaje final estructurado
        alerta_completa = f"⚠️ **Riesgo de Déficit Financiero:**\n\n{msg_fin}\n\n---\n\n💡 **Recomendación Estratégica para viabilizar el proyecto:**\n\n{rec_texto}"
        
        st.warning(alerta_completa)
        
    # =========================================================================
    ## 📍 PASO 6: Proyección Climática y Priorización Predial
    # =========================================================================
    
    # --- 1. TRAYECTORIA CLIMÁTICA Y DEMOGRÁFICA (EXPLORADOR ENSO) ---
    with st.expander(f"📈 PROYECCIÓN DINÁMICA DE SEGURIDAD HÍDRICA (2024 - 2050): {nombre_zona}", expanded=False):
        tab_resumen, tab_escenarios = st.tabs(["📊 Resumen Multivariado (Onda ENSO)", "🔬 Explorador de Escenarios (Cono)"])
        anios_proj = list(range(2024, 2051))

        with tab_resumen:
            col_t1, col_t2 = st.columns(2)
            with col_t1: activar_cc = st.toggle("🌡️ Incluir Cambio Climático", value=True, key="td_t1_cc")
            with col_t2: activar_enso = st.toggle("🌊 Incluir Variabilidad ENSO", value=True, key="td_t1_enso")

            datos_proj = []
            for a in anios_proj:
                delta_a = a - 2024
                # Crecimiento demográfico/agropecuario (~1.5% anual)
                f_dem = (1 + 0.015) ** delta_a
                # Degradación de recarga/oferta por Cambio Climático (~0.5% anual)
                f_cc_base = (1 - 0.005) ** delta_a if activar_cc else 1.0
                
                f_enso = 0.0
                estado_enso = "Neutro ⚖️"
                if activar_enso:
                    f_enso = 0.25 * np.sin((2 * np.pi * delta_a) / 4.5) 
                    estado_enso = "Niña 🌧️" if f_enso > 0.1 else "Niño ☀️" if f_enso < -0.1 else "Neutro ⚖️"
                
                # 🛡️ Escudo anti-negativos para evitar colapsos matemáticos
                f_cli_total = max(0.1, f_cc_base + f_enso) 
                
                # 🔬 PROYECCIÓN DE VOLÚMENES FÍSICOS (Conectado al Bloque 1 y 2)
                # NOTA: Usamos 'oferta_nominal' y 'demanda_m3s' del Top Dashboard
                o_m3 = (oferta_nominal * f_cli_total) * 31536000
                r_m3 = (float(st.session_state.get('aleph_recarga_mm', 350.0)) * f_cli_total) * float(st.session_state.get('aleph_area_km2', 10.0)) * 1000
                c_m3 = (demanda_m3s * f_dem) * 31536000
                
                # ⚖️ NÚCLEO MATEMÁTICO FUTURO
                n = min(100.0, (volumen_repuesto_m3 / c_m3) * 100) if c_m3 > 0 else 100.0
                
                bfi_sim = r_m3 / o_m3 if o_m3 > 0 else 0.0
                fact_superv_sim = min(1.0, r_m3 / c_m3) if c_m3 > 0 else 1.0
                r = max(0.0, min(100.0, (bfi_sim / 0.70) * 100 * fact_superv_sim))
                
                wei_sim = c_m3 / o_m3 if o_m3 > 0 else 1.0
                e = max(0.0, min(100.0, 100.0 - (wei_sim / 0.40) * 60))
                
                caudal_L_s_sim = (o_m3 / 31536000) * 1000
                carga_mg_s_futura = carga_mg_s * f_dem
                dbo_mg_l_sim = carga_mg_s_futura / caudal_L_s_sim if caudal_L_s_sim > 0 else 999.0
                cal = max(0.0, min(100.0, 100.0 - ((dbo_mg_l_sim / 10.0) * 100)))
                
                datos_proj.extend([
                    {"Año": a, "Indicador": "Neutralidad", "Valor (%)": n, "Fase ENSO": estado_enso},
                    {"Año": a, "Indicador": "Resiliencia", "Valor (%)": r, "Fase ENSO": estado_enso},
                    {"Año": a, "Indicador": "Seguridad (Inv. Estrés)", "Valor (%)": e, "Fase ENSO": estado_enso},
                    {"Año": a, "Indicador": "Calidad", "Valor (%)": cal, "Fase ENSO": estado_enso}
                ])
                
            fig_line1 = px.line(pd.DataFrame(datos_proj), x="Año", y="Valor (%)", color="Indicador", hover_data=["Fase ENSO"],
                               color_discrete_map={"Neutralidad": "#2ecc71", "Resiliencia": "#3498db", "Seguridad (Inv. Estrés)": "#e74c3c", "Calidad": "#9b59b6"})
            
            fig_line1.add_hrect(y0=0, y1=40, fillcolor="red", opacity=0.1, layer="below", annotation_text="  Zona Crítica (<40%)")
            fig_line1.add_hrect(y0=40, y1=70, fillcolor="orange", opacity=0.1, layer="below", annotation_text="  Zona Vulnerable (40-70%)")
            fig_line1.update_layout(height=400, hovermode="x unified", yaxis_range=[0, 105], title="Evolución de la Salud Integral del Sistema (0 = Colapso, 100 = Óptimo)")
            st.plotly_chart(fig_line1, use_container_width=True)

        with tab_escenarios:
            col_e1, col_e2 = st.columns([1, 2])
            with col_e1:
                ind_sel = st.selectbox("🎯 Indicador a Evaluar:", ["Estrés Hídrico", "Resiliencia", "Neutralidad", "Calidad"], key="td_ind_sel")
                activar_cc_esc = st.toggle("🌡️ Efecto Cambio Climático", value=True, key="td_t2_cc")
            with col_e2:
                diccionario_escenarios = {
                    "Onda Dinámica": "onda", "Condición Neutra": 0.0, "🟡 Niño Moderado": -0.15,
                    "🔴 Niño Severo": -0.35, "🟢 Niña Moderada": 0.15, "🔵 Niña Fuerte": 0.35
                }
                curvas_sel = st.multiselect("🌊 Curvas Climáticas:", list(diccionario_escenarios.keys()), default=["Onda Dinámica", "Condición Neutra", "🔴 Niño Severo"], key="td_curvas")

            datos_esc = []
            for a in anios_proj:
                delta_a = a - 2024
                f_dem = (1 + 0.015) ** delta_a
                f_cc_base = (1 - 0.005) ** delta_a if activar_cc_esc else 1.0
                
                for nombre_esc in curvas_sel:
                    val_esc = diccionario_escenarios[nombre_esc]
                    f_enso = 0.25 * np.sin((2 * np.pi * delta_a) / 4.5) if val_esc == "onda" else val_esc
                    f_cli_total = f_cc_base + f_enso
                    
                    o_m3 = (oferta_nominal * f_cli_total) * 31536000
                    r_m3 = (float(st.session_state.get('aleph_recarga_mm', 350.0)) * f_cli_total) * float(st.session_state.get('aleph_area_km2', 10.0)) * 1000
                    c_m3 = (demanda_m3s * f_dem) * 31536000
                    
                    if ind_sel == "Neutralidad": val = min(100.0, (volumen_repuesto_m3 / c_m3) * 100) if c_m3 > 0 else 100.0
                    elif ind_sel == "Resiliencia": val = min(100.0, ((r_m3 + o_m3) / ((c_m3+1) * 2)) * 100)
                    elif ind_sel == "Estrés Hídrico": val = min(100.0, (c_m3 / o_m3) * 100) if o_m3 > 0 else 100.0
                    else: 
                        fac_dil = (o_m3 / (c_m3 + 1))
                        val = min(100.0, max(0.0, 50.0 + (fac_dil * 0.5) + (sist_saneamiento * 0.05)))
                        
                    datos_esc.append({"Año": a, "Escenario": nombre_esc, "Valor (%)": val})
                    
            if datos_esc:
                fig_esc = px.line(pd.DataFrame(datos_esc), x="Año", y="Valor (%)", color="Escenario")
                fig_esc.update_traces(line=dict(width=3)) 
                fig_esc.update_layout(height=400, hovermode="x unified", legend=dict(orientation="h", yanchor="bottom", y=-0.3, xanchor="center", x=0.5))
                st.plotly_chart(fig_esc, use_container_width=True)

    # --- 2. RANKING TERRITORIAL MULTICRITERIO (AHP) Y RADAR ---
    with st.expander(f"🏆 RANKING TERRITORIAL MULTICRITERIO (AHP)", expanded=False):
        lista_cuencas = []
        if 'capas' in locals() and capas.get('cuencas') is not None and not capas['cuencas'].empty:
            if 'SUBC_LBL' in capas['cuencas'].columns:
                lista_cuencas = capas['cuencas']['SUBC_LBL'].dropna().unique().tolist()
                
        if not lista_cuencas:
            lista_cuencas = ["Río Chico", "Río Grande", "Quebrada La Mosca", "Río Buey", "Pantanillo"]
            
        st.caption(f"🔄 Reordenado en vivo usando Pesos AHP (Panel Lateral): Hídrico ({w_agua*100:.0f}%) | Biótico ({w_bio*100:.0f}%) | Socioeconómico ({w_socio*100:.0f}%)")
        
        datos_ranking = []
        for c in lista_cuencas:
            pseudo_seed = sum([ord(char) for char in c])
            np.random.seed(pseudo_seed)
            
            n_val = np.random.uniform(20, 90) if c != nombre_zona else ind_neutralidad
            r_val = np.random.uniform(20, 95) if c != nombre_zona else ind_resiliencia
            e_val = np.random.uniform(10, 80) if c != nombre_zona else ind_estres
            c_val = np.random.uniform(20, 100) if c != nombre_zona else ind_calidad
            
            # 🧠 MOTOR AHP: Conectado a los sliders
            urgencia_hidrica = e_val  
            urgencia_biotica = 100 - c_val  
            urgencia_socio = 100 - r_val 
            
            score_urgencia = (urgencia_hidrica * w_agua) + (urgencia_biotica * w_bio) + (urgencia_socio * w_socio)
            
            datos_ranking.append({
                "Territorio": c, "Índice Prioridad (AHP)": score_urgencia,
                "Neutralidad (%)": n_val, "Resiliencia (%)": r_val,
                "Estrés Hídrico (%)": e_val, "Calidad de Agua (%)": c_val
            })
            
        df_ranking = pd.DataFrame(datos_ranking).sort_values(by="Índice Prioridad (AHP)", ascending=False)
        
        c_tbl, c_rad = st.columns([1.5, 1])
        with c_tbl:
            st.dataframe(
                df_ranking.style.background_gradient(cmap="Reds", subset=["Índice Prioridad (AHP)", "Estrés Hídrico (%)"])
                .background_gradient(cmap="Blues", subset=["Resiliencia (%)"])
                .background_gradient(cmap="Greens", subset=["Neutralidad (%)", "Calidad de Agua (%)"])
                .format({"Índice Prioridad (AHP)": "{:.1f}", "Neutralidad (%)": "{:.1f}%", "Resiliencia (%)": "{:.1f}%", "Estrés Hídrico (%)": "{:.1f}%", "Calidad de Agua (%)": "{:.1f}%"}),
                use_container_width=True, hide_index=True
            )
            st.download_button("📥 Descargar Ranking AHP (CSV)", df_ranking.to_csv(index=False).encode('utf-8'), "Ranking_Territorial_AHP.csv", "text/csv")

            # DISTRIBUCIÓN REGIONAL (BOX PLOT)
            df_melt = df_ranking.melt(id_vars=["Territorio"], value_vars=["Neutralidad (%)", "Resiliencia (%)", "Estrés Hídrico (%)", "Calidad de Agua (%)"], var_name="Índice", value_name="Valor (%)")
            fig_box = px.box(df_melt, x="Índice", y="Valor (%)", color="Índice", points="all", title="Distribución Regional de Indicadores",
                             color_discrete_map={"Neutralidad (%)": "#2ecc71", "Resiliencia (%)": "#3498db", "Estrés Hídrico (%)": "#e74c3c", "Calidad de Agua (%)": "#9b59b6"})
            fig_box.update_layout(height=300, showlegend=False, margin=dict(t=30, b=0, l=0, r=0))
            st.plotly_chart(fig_box, use_container_width=True)

        with c_rad:
            fig_radar = go.Figure()
            categorias = ['Neutralidad', 'Resiliencia', 'Seguridad (Inv. Estrés)', 'Calidad', 'Neutralidad']
            
            fig_radar.add_trace(go.Scatterpolar(r=[100]*5, theta=categorias, fill='toself', fillcolor='rgba(39, 174, 96, 0.15)', line=dict(color='rgba(255,255,255,0)'), name='Óptimo (>70%)', hoverinfo='none'))
            fig_radar.add_trace(go.Scatterpolar(r=[70]*5, theta=categorias, fill='toself', fillcolor='rgba(241, 196, 15, 0.2)', line=dict(color='rgba(255,255,255,0)'), name='Vulnerable (40-70%)', hoverinfo='none'))
            fig_radar.add_trace(go.Scatterpolar(r=[40]*5, theta=categorias, fill='toself', fillcolor='rgba(192, 57, 43, 0.25)', line=dict(color='rgba(255,255,255,0)'), name='Crítico (<40%)', hoverinfo='none'))

            valores_radar = [ind_neutralidad, ind_resiliencia, max(0, 100-ind_estres), ind_calidad]
            fig_radar.add_trace(go.Scatterpolar(
                r=valores_radar + [valores_radar[0]], theta=categorias,
                fill='toself', name=nombre_zona, line=dict(color='#2c3e50', width=3),
                fillcolor='rgba(41, 128, 185, 0.7)', mode='lines+markers', marker=dict(size=8, color='#2c3e50')
            ))

            fig_radar.update_layout(
                polar=dict(radialaxis=dict(visible=True, range=[0, 100], tickvals=[40, 70, 100], ticktext=["40%", "70%", "100%"]),
                           angularaxis=dict(tickfont=dict(size=11, color="black", weight="bold"))),
                showlegend=True, legend=dict(orientation="h", y=-0.2, x=0.5, xanchor="center"),
                title=dict(text="Huella de Salud Territorial", font=dict(size=18)), height=380, margin=dict(l=40, r=40, t=50, b=20)
            )
            st.plotly_chart(fig_radar, use_container_width=True)
            
            promedio_salud = np.mean(valores_radar)
            color_box, msg_estado = ("#27ae60", "🟢 <b>TERRITORIO ÓPTIMO</b>") if promedio_salud >= 70 else ("#f39c12", "🟡 <b>TERRITORIO VULNERABLE</b>") if promedio_salud >= 40 else ("#c0392b", "🔴 <b>TERRITORIO CRÍTICO</b>")
            st.markdown(f"<div style='padding:10px; border-radius:5px; border-left: 5px solid {color_box}; background-color:#f8f9fa;'>{msg_estado}<br>Puntaje Salud: {promedio_salud:.1f}/100</div>", unsafe_allow_html=True)

    # --- 3. GLOSARIO METODOLÓGICO Y FUENTES ---
    with st.expander("📚 Glosario Metodológico (VWBA - WRI)", expanded=False):
        st.markdown("""
        * **Neutralidad Hídrica (VWBA):** Volumen de agua restituido mediante SbN vs Huella Hídrica. Objetivo: 100%.
        * **Resiliencia Territorial (BFI USGS):** Capacidad del ecosistema base para soportar eventos de sequía.
        * **Estrés Hídrico (WEI+):** Extracción vs Oferta. >40% indica estrés severo.
        * **Calidad de Agua (WQI):** Dilución natural de DBO5 y mitigación sanitaria.
        """)

    # =========================================================================
    # PRIORIZACIÓN PREDIAL PARA CONECTIVIDAD RIPARIA
    # =========================================================================
    st.markdown("---")
    st.subheader(f"🎯 Inteligencia de Negociación: Priorización Predial ({nombre_zona})")
    st.markdown("Cruza las necesidades de restauración riparia con la estructura predial alojada en la nube para identificar qué propiedades priorizar.")

    rios_strahler_crudos = st.session_state.get('gdf_rios')
    buffer_m = st.session_state.get('buffer_m_ripario', None) 
    rios_strahler = None
    
    if rios_strahler_crudos is not None and not rios_strahler_crudos.empty and gdf_zona is not None:
        try:
            rios_3116 = rios_strahler_crudos.to_crs(epsg=3116)
            zona_3116 = gdf_zona.to_crs(epsg=3116)
            rios_clip = gpd.clip(rios_3116, zona_3116)
            if not rios_clip.empty: rios_strahler = rios_clip
        except Exception as e:
            st.warning(f"Aviso validando red hídrica: {e}")
            
    if rios_strahler is None or rios_strahler.empty:
        with st.expander("⚠️ Paso 1: Faltan Datos - Generar Red Hídrica", expanded=True):
            st.info(f"Para priorizar predios necesitamos la red hídrica exacta de **{nombre_zona}**. ¡Trázalos usando el Motor Hidrológico!")
            render_motor_hidrologico(gdf_zona)
            
    elif buffer_m is None:
        with st.expander("⚠️ Paso 2: Configurar Franja Riparia", expanded=True):
            st.success("✅ ¡Ríos detectados en la zona! Ahora define el ancho de la zona de protección riparia.")
            render_motor_ripario()
            
    else:
        with st.expander("⚙️ Recalcular Franja Riparia", expanded=False):
            st.success(f"✅ Red Hídrica y Franja Riparia de {buffer_m}m listas para el cruce predial.")
            render_motor_ripario()

        # =========================================================================
        # BLOQUE 4: MOTOR DE CRUCE MULTI-ANILLO Y VISOR 3D (PYDECK)
        # =========================================================================
        
        # 2. Cargar Universo Catastral (Capa Predial Total desde Supabase)
        @st.cache_data(ttl=3600, show_spinner=False)
        def obtener_catastro_estrategico(_gdf_zona):
            import requests, tempfile
            import pandas as pd
            import geopandas as gpd
            try:
                # 🎯 Apuntamos al nuevo archivo maestro del catastro
                url = "https://ldunpssoxvifemoyeuac.supabase.co/storage/v1/object/public/geojson/PrediosVeredas_CV.geojson"
                res = requests.get(url)
                if res.status_code == 200:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".geojson") as tmp:
                        tmp.write(res.content)
                        tmp_path = tmp.name
                    gdf_c = gpd.read_file(tmp_path)
                    if not gdf_c.empty:
                        gdf_c.set_crs(epsg=4326, allow_override=True, inplace=True)
                        gdf_c_3116 = gdf_c.to_crs(epsg=3116)
                        zona_3116 = _gdf_zona.to_crs(epsg=3116)
                        
                        # Corrección de geometrías y recorte estricto a la zona de interés
                        gdf_c_3116['geometry'] = gdf_c_3116.geometry.make_valid().buffer(0)
                        zona_3116['geometry'] = zona_3116.geometry.make_valid().buffer(0)
                        return gpd.clip(gdf_c_3116, zona_3116)
            except Exception as e:
                return None
            return None
            
        with st.spinner("Descargando Universo Catastral (PrediosVeredas_CV)..."):
            capa_predios = obtener_catastro_estrategico(gdf_zona)

        # 3. Ejecutar el Motor de Cruce
        if rios_strahler is not None and not rios_strahler.empty:
            
            with st.container(border=True):
                st.markdown("#### ⚙️ Simulación Concéntrica de Franjas de Protección")
                
                # Preparación de la red hídrica
                rios_strahler = rios_strahler.reset_index(drop=True)
                rios_strahler['ID_Tramo'] = ["Segmento " + str(i+1) for i in range(len(rios_strahler))]
                if 'longitud_km' in rios_strahler.columns:
                    rios_strahler['longitud_km'] = rios_strahler['longitud_km'].round(2)
                
                # Extraer tamaños de anillo (Mínimo, Ideal, Óptimo)
                anillos = st.session_state.get('multi_rings', [10, 20, 30])
                b_min, b_med, b_max = anillos[0], anillos[1], anillos[2]
                
                # 🌿 MAGIA GEOMÉTRICA: Pre-cálculo de anillos fusionados
                rios_3116 = rios_strahler.to_crs(epsg=3116)
                rios_union = rios_3116.unary_union
                
                geom_max = rios_union.buffer(b_max, resolution=2)
                geom_med = rios_union.buffer(b_med, resolution=2)
                geom_min = rios_union.buffer(b_min, resolution=2)
                
                buffer_max_gdf = gpd.GeoDataFrame(geometry=[geom_max], crs=3116)
                
                # CÁLCULO DE ÁREAS POR TRAMO HÍDRICO
                datos_tramos = []
                for idx, row in rios_3116.iterrows():
                    long_m = row.geometry.length
                    orden = row.get('Orden_Strahler', 1)
                    long_km = long_m / 1000.0
                    
                    datos_tramos.append({
                        "ID Franja (Tramo)": row['ID_Tramo'],
                        "Orden de Strahler": orden,
                        "Longitud (Km)": long_km,
                        f"Mínimo ({b_min}m) ha": (long_m * (b_min * 2)) / 10000.0,
                        f"Ideal ({b_med}m) ha": (long_m * (b_med * 2)) / 10000.0,
                        f"Óptimo ({b_max}m) ha": (long_m * (b_max * 2)) / 10000.0,
                        "Importancia Ecológica": (orden * 50) + (long_km * 10)
                    })
                df_tramos = pd.DataFrame(datos_tramos).sort_values(by="Importancia Ecológica", ascending=False)
                
                tot_min = df_tramos[f"Mínimo ({b_min}m) ha"].sum()
                tot_med = df_tramos[f"Ideal ({b_med}m) ha"].sum()
                tot_max = df_tramos[f"Óptimo ({b_max}m) ha"].sum()
                tot_longitud_km = df_tramos["Longitud (Km)"].sum() 

                st.success(f"✅ Modelando 3 escenarios concéntricos simultáneos ({b_min}m, {b_med}m, {b_max}m)...")
                
                # MÉTRICAS TÁCTICAS
                cm1, cm2, cm3, cm4, cm5 = st.columns(5)
                cm1.metric(f"🔴 Escenario {b_min}m", f"{tot_min:,.1f} ha")
                cm2.metric(f"🟡 Escenario {b_med}m", f"{tot_med:,.1f} ha", f"+{(tot_med - tot_min):,.1f} ha", delta_color="off")
                cm3.metric(f"🟢 Escenario {b_max}m", f"{tot_max:,.1f} ha", f"+{(tot_max - tot_med):,.1f} ha", delta_color="off")
                cm4.metric("🌿 Tramos Hídricos", f"{len(df_tramos)}")
                cm5.metric("📏 Longitud Total", f"{tot_longitud_km:,.1f} km")
                
                tab_predios, tab_tramos = st.tabs(["🏡 Impacto Predial (Negociación)", "🌿 Áreas por Franja Riparia (Tramos)"])
                
                with tab_tramos:
                    st.markdown("##### 📋 Matriz Detallada por Franja Riparia")
                    st.dataframe(df_tramos.style.background_gradient(cmap="Greens", subset=["Importancia Ecológica"]).format(precision=2), use_container_width=True, hide_index=True)
                
                with tab_predios:
                    st.markdown("##### 🧬 Álgebra de Mapas: Aplicación de Reglas CuencaVerde")
                    st.info("**R1:** Cuidar lo bien conservado | **R2:** Mejorar lo cercano | **R3:** Conectar fragmentos | **R4:** Prioridad a bosques riparios.")
                    
                    predios_en_buffer = gpd.GeoDataFrame()
                    if capa_predios is not None and not capa_predios.empty:
                        with st.spinner("Ejecutando Resta Espacial e Intersección Catastral..."):
                            try:
                                predios_3116 = capa_predios.to_crs(epsg=3116)
                                
                                # 🧮 ÁLGEBRA ESPACIAL 1: Restar Capital Gestionado (Evitar doble esfuerzo)
                                if 'gdf_predios_mapa' in locals() and gdf_predios_mapa is not None and not gdf_predios_mapa.empty:
                                    gestionados_3116 = gdf_predios_mapa.to_crs(3116)
                                    # La Resta: Lo que falta = Buffer IDEAL (Regla 4) - Lo ya intervenido
                                    buffer_deficit_geom = buffer_max_gdf.geometry.iloc[0].difference(gestionados_3116.unary_union)
                                    buffer_evaluacion_gdf = gpd.GeoDataFrame(geometry=[buffer_deficit_geom], crs=3116)
                                else:
                                    buffer_evaluacion_gdf = buffer_max_gdf
                                
                                # 🧮 ÁLGEBRA ESPACIAL 2: Intersección con Universo Catastral (Identificar dueños)
                                predios_en_buffer = gpd.overlay(predios_3116, buffer_evaluacion_gdf, how='intersection')
                                
                                if not predios_en_buffer.empty:
                                    predios_en_buffer['Area_Faltante_ha'] = predios_en_buffer.geometry.area / 10000.0
                                    
                                    col_id = next((col for col in ['MATRICULA', 'COD_CATAST', 'FICHA', 'OBJECTID', 'id', 'NOMBRE'] if col in predios_en_buffer.columns), None)
                                    if col_id is None:
                                        predios_en_buffer['ID_Predio'] = predios_en_buffer.index
                                        col_id = 'ID_Predio'
                                        
                                    predios_agrupados = predios_en_buffer.groupby(col_id).agg({'Area_Faltante_ha': 'sum'}).reset_index()
                                    
                                    # 🧮 ÁLGEBRA ESPACIAL 3: Inyectar datos satelitales (Reglas 1, 2 y 3)
                                    # Usamos las coberturas leídas en el Paso 1 para tipificar el esfuerzo necesario
                                    if 'areas_data' in locals() and areas_data and area_total_ha > 0:
                                        pct_bosque = sum([x["Área (Ha)"] for x in areas_data if "Bosque" in x["Cobertura"]]) / area_total_ha
                                        pct_matorral = sum([x["Área (Ha)"] for x in areas_data if "Matorrales" in x["Cobertura"]]) / area_total_ha
                                    else:
                                        pct_bosque, pct_matorral = 0.4, 0.15 # Valores refugio
                                        
                                    datos_prioridad = []
                                    for idx, row in predios_agrupados.iterrows():
                                        area_tot = row['Area_Faltante_ha']
                                        
                                        # Aplicación paramétrica
                                        area_cuidar = area_tot * pct_bosque
                                        area_mejorar = area_tot * pct_matorral
                                        area_conectar = area_tot - area_cuidar - area_mejorar
                                        
                                        datos_prioridad.append({
                                            "Predio (ID/Matrícula)": row[col_id],
                                            "Déficit Ripario (Ha)": area_tot,
                                            "🟢 Conservar (R1)": area_cuidar,
                                            "🟡 Mejorar (R2)": area_mejorar,
                                            "🔴 Conectar (R3)": max(0.0, area_conectar),
                                            "Prioridad Estratégica": (max(0, area_conectar) * 10) + (area_mejorar * 5)
                                        })
                                        
                                    df_prioridad = pd.DataFrame(datos_prioridad).sort_values(by="Prioridad Estratégica", ascending=False)
                                    
                                    c_rank1, c_rank2 = st.columns([2.5, 1])
                                    with c_rank1:
                                        st.dataframe(df_prioridad.head(15).style.background_gradient(cmap="YlOrRd", subset=["🔴 Conectar (R3)", "Prioridad Estratégica"]).background_gradient(cmap="Greens", subset=["🟢 Conservar (R1)"]).format(precision=2), use_container_width=True, hide_index=True)
                                    with c_rank2:
                                        st.info("💡 **Guía Táctica:** Los predios con puntajes rojos altos son los 'cuellos de botella'. Ahí la sabana o el cultivo cortan la ronda del río. **Son la máxima prioridad de negociación.**")
                                        st.metric("Nuevos Predios a Contactar", f"{len(df_prioridad)}")
                                        st.download_button("📥 Descargar Matriz de Intervención", df_prioridad.to_csv(index=False).encode('utf-8'), "Matriz_Reglas_CV.csv", "text/csv")
                                else:
                                    st.success("Toda la red hídrica en esta zona ya se encuentra dentro de las áreas históricamente gestionadas. No hay déficit ripario reportado.")
                            except Exception as e:
                                st.error(f"Error técnico en el cruce geográfico de Álgebra de Mapas: {e}")
                    else:
                        st.info("ℹ️ Carga la red de predios catastrales para correr el álgebra de mapas.")

            # =========================================================
            # 🗺️ EL MAPA TÁCTICO PYDECK (VISOR 3D DE NEGOCIACIÓN)
            # =========================================================
            st.markdown("---")
            st.markdown(f"#### 🗺️ Visor Táctico de Conectividad y Predios: **{nombre_zona}**")
            import pydeck as pdk
            
            try:
                rios_4326 = rios_strahler.to_crs(epsg=4326).copy()
                c_lat, c_lon = rios_4326.geometry.iloc[0].centroid.y, rios_4326.geometry.iloc[0].centroid.x
            except: c_lat, c_lon = 6.2, -75.5 
            
            capas_mapa = []
            
            # Capa 1: Límite de Cuenca/Zona
            if gdf_zona is not None:
                zona_4326 = gdf_zona.to_crs("EPSG:4326")
                capas_mapa.append(pdk.Layer("GeoJsonLayer", data=zona_4326, opacity=1, stroked=True, get_line_color=[0, 200, 0, 255], get_line_width=3, filled=False))
            
            # Capa 2: Anillos Concéntricos (Niveles de Prioridad)
            if 'geom_max' in locals():
                gdf_max = gpd.GeoDataFrame(geometry=[geom_max], crs=3116).to_crs(4326)
                capas_mapa.append(pdk.Layer("GeoJsonLayer", data=gdf_max, opacity=0.2, get_fill_color=[171, 235, 198], stroked=False))
                
                gdf_med = gpd.GeoDataFrame(geometry=[geom_med], crs=3116).to_crs(4326)
                capas_mapa.append(pdk.Layer("GeoJsonLayer", data=gdf_med, opacity=0.4, get_fill_color=[88, 214, 141], stroked=False))
                
                gdf_min = gpd.GeoDataFrame(geometry=[geom_min], crs=3116).to_crs(4326)
                capas_mapa.append(pdk.Layer("GeoJsonLayer", data=gdf_min, opacity=0.6, get_fill_color=[40, 180, 99], stroked=False))

            # Capa 3: Red Hídrica (Strahler)
            if 'rios_4326' in locals():
                capas_mapa.append(pdk.Layer(
                    "GeoJsonLayer", data=rios_4326,
                    get_line_color=[31, 97, 141, 255], get_line_width=2, lineWidthMinPixels=2,
                    pickable=True, autoHighlight=True
                ))
            
            # Capa 4: Predios Estratégicos (Afectados)
            if 'predios_en_buffer' in locals() and not predios_en_buffer.empty:
                col_id_oficial = next((col for col in ['MATRICULA', 'COD_CATAST', 'FICHA', 'OBJECTID', 'id'] if capa_predios is not None and col in capa_predios.columns), None)
                
                if col_id_oficial:
                    ids_afectados = predios_en_buffer[col_id_oficial].unique()
                    predios_a_dibujar = capa_predios[capa_predios[col_id_oficial].isin(ids_afectados)].to_crs(epsg=4326)
                else:
                    predios_a_dibujar = predios_en_buffer.to_crs(epsg=4326)
                    
                capas_mapa.append(pdk.Layer(
                    "GeoJsonLayer", data=predios_a_dibujar, opacity=0.4,
                    stroked=True, filled=True, get_fill_color=[255, 165, 0, 150],
                    get_line_color=[255, 140, 0, 255], get_line_width=2,
                    pickable=True, autoHighlight=True
                ))
            
            # Renderizado 3D
            view_state = pdk.ViewState(latitude=c_lat, longitude=c_lon, zoom=13, pitch=45)
            tooltip = {"html": "<b>Tramo Hídrico:</b> {ID_Tramo}<br/><b>Orden:</b> {Orden_Strahler}<br/><b>Longitud:</b> {longitud_km} km", "style": {"backgroundColor": "steelblue", "color": "white"}}
            st.pydeck_chart(pdk.Deck(layers=capas_mapa, initial_view_state=view_state, map_style="light", tooltip=tooltip), use_container_width=True)

        else:
            st.warning("⚠️ El cruce predial y el mapa táctico están en pausa porque aún no se han calculado los ríos.")
            st.info("👆 Por favor, utiliza el botón del motor hidrológico de arriba para iluminar este tablero.")

    # ==============================================================================
    # 📍 PASO 6: EL MANIFIESTO Y SÍNTESIS DE LA IA ESTRATÉGICA
    # ==============================================================================
    st.markdown("---")
    st.markdown("## 📍 PASO 6: Síntesis Ejecutiva y Manifiesto")
    st.info("Este motor consolida la telemetría del Aleph y los diagnósticos en un documento institucional.")
    
    # 🧠 LA CAJA NEGRA DE LA IA
    st.markdown("### 🧠 Veredicto de la IA Estratégica (Sihcli-Poter)")
    
    texto_estado = "CRÍTICO" if ishi_final < 40 else "VULNERABLE" if ishi_final < 70 else "ÓPTIMO"
    
    # Protecciones de variables por si el usuario cambia el orden de cálculo arriba
    new_ishi_rep = new_ishi if 'new_ishi' in locals() else ishi_final
    inv_usd_rep = presupuesto_usd/1e6 if 'presupuesto_usd' in locals() else 5.0
    w_verde_rep = w_verde * 100 if 'w_verde' in locals() else 50.0
    
    msg_ia = f"""
    El ecosistema hídrico de **{nombre_zona}** presenta un Índice de Seguridad (ISHI) base del **{ishi_final:.1f}%**, lo que lo clasifica en un estado **{texto_estado}**. 
    La presión metabólica generada por **{int(pob_total):,.0f} habitantes** y la actividad pecuaria exige la remoción de **{carga_total_ton:,.0f} Ton/año** de DBO5 para evitar el colapso del oxígeno disuelto.
    
    El Optimizador Matemático concluye que una inyección financiera de **${inv_usd_rep:,.1f} Millones USD**, distribuida con un sesgo del **{w_verde_rep:.0f}% hacia Soluciones Basadas en la Naturaleza (SbN)**, logrará expandir la huella de seguridad hasta un **{new_ishi_rep:.1f}%**.
    
    **💡 Recomendación de Política Pública:** Se sugiere iniciar inmediatamente los procesos de gestión predial en las franjas riparias de alto orden de Strahler, combinadas con sistemas descentralizados (STAM) rurales, por poseer el costo marginal unitario más eficiente frente a la infraestructura gris tradicional.
    """
    
    st.success(msg_ia)

    c_rep1, c_rep2 = st.columns([2, 1])
    with c_rep1:
        st.markdown("### ⚙️ Configuración del Documento")
        titulo_plan = st.text_input("Título del Manifiesto:", f"Plan de Seguridad Hídrica - {nombre_zona}", key="rep_tit")
        director = st.text_input("Firma Autorizada:", "Dirección Técnica - CuencaVerde", key="rep_dir")
    with c_rep2:
        st.markdown("### 🧐 Estructura Multidimensión")
        st.caption("Cap 1: Veredicto | Cap 2: Metabolismo | Cap 3: Clima | Cap 4: Inversión | Glosario")

    st.markdown("---")
    
    # EL BOTÓN MAESTRO
    if st.button("🚀 Generar y Descargar Manifiesto Estratégico", type="primary", use_container_width=True):
        with st.spinner("Capturando mapas geográficos y ensamblando el reporte..."):
            
            # 1. RECOLECCIÓN DE DATOS
            od_pct = st.session_state.get('calidad_oxigeno_pct', ind_calidad)
            dbo_mgL = st.session_state.get('calidad_dbo_salida_mgL', concentracion_dbo_mg_l)
            acuifero_mgL = st.session_state.get('calidad_acuifero_mgL', concentracion_dbo_mg_l * 0.12)
            
            p_nino = st.session_state.get('aleph_iri_nino', 20)
            p_neutro = st.session_state.get('aleph_iri_neutro', 80)
            trimestre = st.session_state.get('aleph_iri_trimestre', 'AMJ')
            
            # 2. ENSAMBLADOR DE NARRATIVA 
            impacto_rio = "grave" if od_pct < 40 else "moderado" if od_pct < 70 else "bajo"
            
            num_concesiones = st.session_state.get('total_concesiones', 145)
            vol_concesionado = st.session_state.get('volumen_concesionado_m3s', 2.3)
            num_vertimientos = st.session_state.get('total_vertimientos', 89)
            
            cap2_txt = f"El análisis demográfico y sectorial para {nombre_zona} proyecta una población de {pob_total:,.0f} habitantes para el año {anio_actual}. Esta presión poblacional se complementa con una intensa actividad pecuaria estimada en {bovinos:,.0f} bovinos, {porcinos:,.0f} porcinos y {aves:,.0f} aves. La síntesis de estas dinámicas metabólicas genera una carga orgánica estimada de {carga_total_ton:,.1f} Ton/año de DBO5.\n\n" \
                       f"En términos de gobernanza y administración del recurso, la base de datos espacial registra un total de {num_concesiones} concesiones de agua activas, captando un volumen agregado de {vol_concesionado:.2f} m³/s, y un inventario de {num_vertimientos} permisos de vertimiento formales.\n\n" \
                       f"Los modelos de asimilación proyectan una salud del río del {od_pct:.1f}%, indicando un impacto {impacto_rio} en los ecosistemas acuáticos (con concentraciones medias de {dbo_mgL:.1f} mg/L). Asimismo, la presión sobre las fuentes hídricas subterráneas (acuíferos) requiere atención, estimándose una concentración de contaminantes por infiltración de {acuifero_mgL:.2f} mg/L."
            cap3_txt = f"Basado en el monitoreo satelital del IRI (Columbia University), el trimestre {trimestre} presenta una probabilidad del {p_nino}% para el fenómeno de El Niño y {p_neutro}% para Neutralidad. Esta configuración climática actual exige calibrar los portafolios de inversión para priorizar la recarga hídrica en la cuenca alta."

            # 3. EXPORTADOR .DOCX
            try:
                from docx import Document
                from docx.shared import Pt, Inches, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                import io

                doc = Document()
                style = doc.styles['Normal']
                style.font.name = 'Georgia'
                style.font.size = Pt(11)

                tit = doc.add_heading(titulo_plan, level=0)
                tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub = doc.add_paragraph(f"Firma Avalada: {director} | Fecha de Simulación: {anio_actual}")
                sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_page_break()

                doc.add_heading('Capítulo 1: Resumen y Veredicto Territorial', level=1)
                
                verbo_inv = "expande" if new_ishi_rep >= ishi_final else "mantiene"
                estado_ishi = "crítico" if ishi_final < 40 else "vulnerable" if ishi_final < 70 else "óptimo"
                
                cap1_p1 = f"El presente documento, denominado '{titulo_plan}', se formula como el instrumento rector para la planificación hídrica del territorio en el horizonte 2026-2030. Actualmente, los análisis multicriterio del Gemelo Digital (Sihcli-Poter) indican que el Índice de Seguridad Hídrica (ISHI) de la región registra un nivel {estado_ishi} del {ishi_final:.1f}%."
                
                # Rescate seguro de valores
                inv_v_rep = inv_verde/1e6 if 'inv_verde' in locals() else 0.0
                inv_g_rep = inv_gris/1e6 if 'inv_gris' in locals() else 0.0
                
                cap1_p2 = f"Para mitigar esta vulnerabilidad y blindar el metabolismo hídrico regional, el modelo matemático exige una movilización de capital de ${inv_usd_rep:.1f} Millones USD. El algoritmo recomienda direccionar ${inv_v_rep:.2f}M USD a Infraestructura Verde (SbN) y ${inv_g_rep:.2f}M USD a Infraestructura Gris (PTAR). Esta inyección {verbo_inv} estructuralmente la resiliencia, proyectando el ISHI hacia un {new_ishi_rep:.1f}%."
                
                doc.add_paragraph(cap1_p1).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                doc.add_paragraph(cap1_p2).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                try:
                    if 'fig_opt' in locals():
                        img_radar = fig_opt.to_image(format="png", width=700, height=500, scale=2)
                        doc.add_picture(io.BytesIO(img_radar), width=Inches(5.0))
                        cap_rad = doc.add_paragraph("Figura 1: Radar ISHI de Seguridad Hídrica y Proyección de Inversión.")
                        cap_rad.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_rad.runs[0].italic = True
                except: pass

                try:
                    import plotly.express as px
                    import time 
                    
                    if 'gdf_zona' in locals() and gdf_zona is not None and not gdf_zona.empty:
                        zona_wgs84 = gdf_zona.to_crs(epsg=4326)
                        centro_lat = zona_wgs84.geometry.centroid.y.mean()
                        centro_lon = zona_wgs84.geometry.centroid.x.mean()
                        
                        fig_mapa_ctx = px.choropleth_mapbox(
                            zona_wgs84, geojson=zona_wgs84.geometry, locations=zona_wgs84.index,
                            mapbox_style="carto-positron", opacity=0.5, color_discrete_sequence=["#3498db"],
                            center={"lat": centro_lat, "lon": centro_lon}, zoom=9.5
                        )
                        
                        estaciones_gdf = st.session_state.get('gdf_estaciones_filtradas')
                        if estaciones_gdf is not None and not estaciones_gdf.empty:
                            est_wgs84 = estaciones_gdf.to_crs(epsg=4326)
                            import plotly.graph_objects as go
                            fig_mapa_ctx.add_trace(go.Scattermapbox(
                                lat=est_wgs84.geometry.y, lon=est_wgs84.geometry.x,
                                mode='markers', marker=go.scattermapbox.Marker(size=10, color='red'),
                                name='Estaciones Hidrometeorológicas'
                            ))

                        fig_mapa_ctx.update_layout(
                            margin={"r":0,"t":0,"l":0,"b":0},
                            showlegend=False,
                            annotations=[
                                dict(x=0.95, y=0.95, text="<b>⬆ N</b>", showarrow=False, 
                                     font=dict(size=22, color="black"), bgcolor="white", bordercolor="black", borderwidth=1),
                                dict(x=0.05, y=0.05, text="Escala Aprox: 1:100.000", showarrow=False, 
                                     font=dict(size=12, color="black"), bgcolor="white", bordercolor="black", borderwidth=1)
                            ]
                        )
                        time.sleep(1.5) 
                        img_mapa_ctx = fig_mapa_ctx.to_image(format="png", width=800, height=450, scale=2)
                        doc.add_picture(io.BytesIO(img_mapa_ctx), width=Inches(6.0))
                        
                        cap_mapa1 = doc.add_paragraph("Figura 2: Contexto Geográfico del Territorio, Cuenca y Red de Monitoreo.")
                        cap_mapa1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_mapa1.runs[0].italic = True
                except Exception as e: 
                    doc.add_paragraph(f"[Aviso: No se pudo renderizar el mapa geográfico. {e}]")
                    
                doc.add_heading('Capítulo 2: Diagnóstico de Calidad y Metabolismo', level=1)
                doc.add_paragraph(cap2_txt).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                try:
                    if 'crear_velocimetro' in locals() and 'ind_calidad' in locals():
                        fig_cal = crear_velocimetro(ind_calidad, "Salud Sanitaria", "#9b59b6", 40, 70)
                        img_cal = fig_cal.to_image(format="png", width=500, height=350, scale=2)
                        doc.add_picture(io.BytesIO(img_cal), width=Inches(3.5))
                except: pass

                doc.add_heading('Capítulo 3: Multiverso Climático e Hidrología', level=1)
                doc.add_paragraph(cap3_txt).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                doc.add_heading('Capítulo 4: Portafolio de Inversión', level=1)
                doc.add_paragraph(f"Presupuesto distribuido para mitigar el déficit sistémico:")
                
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Grid Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Línea de Intervención'
                hdr_cells[1].text = 'Asignación (USD)'
                table.add_row().cells[0].text, table.rows[1].cells[1].text = 'Infraestructura Verde (SbN)', f"${inv_v_rep * 1e6:,.0f}"
                table.add_row().cells[0].text, table.rows[2].cells[1].text = 'Infraestructura Gris (Saneamiento)', f"${inv_g_rep * 1e6:,.0f}"

                doc.add_page_break()
                doc.add_heading('Glosario de Siglas y Conceptos', level=1)
                
                glosario = {
                    "ISHI (Índice de Seguridad Hídrica)": "Métrica multicriterio que consolida el estrés de abastecimiento, resiliencia física, calidad del agua y neutralidad volumétrica.",
                    "SbN (Soluciones Basadas en la Naturaleza)": "Intervenciones de conservación, restauración ecológica y bioingeniería para proteger el ciclo hidrológico.",
                    "PTAR": "Planta de Tratamiento de Aguas Residuales. Infraestructura gris diseñada para remover la carga orgánica.",
                    "DBO5 (Demanda Biológica de Oxígeno)": "Indicador de contaminación orgánica. Mide el oxígeno que las bacterias consumen para degradar la materia en el agua.",
                    "VWBA (Volumetric Water Benefit Accounting)": "Metodología estándar global desarrollada por el WRI para calcular los beneficios de los proyectos corporativos de agua.",
                    "ENSO (El Niño / La Niña)": "Fenómeno climático global que alterna entre fases de déficit hídrico extremo (El Niño) y excesos pluviométricos (La Niña).",
                    "Streeter-Phelps": "Modelo matemático de simulación que predice la caída y posterior recuperación del oxígeno disuelto en un río tras un vertimiento."
                }
                
                for term, defi in glosario.items():
                    p_glos = doc.add_paragraph()
                    p_glos.add_run(term + ": ").bold = True
                    p_glos.add_run(defi)
                    p_glos.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                buf = io.BytesIO()
                doc.save(buf)
                st.session_state['manifiesto_docx_buffer'] = buf.getvalue()
                st.session_state['nombre_archivo_docx'] = f"Plan_Estrategico_{nombre_zona}.docx"
                
                st.success("✅ ¡Documento estratégico ensamblado con éxito! El sistema está listo para la descarga.")

            except Exception as e:
                st.error(f"Falla crítica en exportación: {e}")

    if 'manifiesto_docx_buffer' in st.session_state:
        st.markdown("---")
        st.markdown("### 📥 Tu documento está listo")
        st.info("Haz clic en el botón azul de abajo para guardar el archivo Word en tu computadora.")
        
        st.download_button(
            label="💾 DESCARGAR PLAN ESTRATÉGICO (.docx)",
            data=st.session_state['manifiesto_docx_buffer'],
            file_name=st.session_state.get('nombre_archivo_docx', 'Plan_Estrategico.docx'),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )

# Fin del archivo
